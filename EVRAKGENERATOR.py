"""
EVRAKGENERATOR - İyileştirilmiş Versiyon
Yazar: Hüseyin İLHAN
Düzenleyen: Claude AI Assistant
"""

import os
import re
import shutil
import datetime
import unicodedata

import pandas as pd
import threading

# tkinter GUI ve belge işlemleri için
import tkinter as tk
from tkinter import messagebox, Toplevel, BooleanVar, Checkbutton
import tkinter.ttk as ttk

# Belge ve GUI için platformlar arası Türkçe karakter destekli font seçimi
import tkinter.font as tkfont_det
_font_root = tk.Tk()
_font_root.withdraw()
_available_fonts = set(tkfont_det.families())
_font_root.destroy()
for _font in ("Arial", "Liberation Sans", "DejaVu Sans", "TkDefaultFont"):
    if _font in _available_fonts:
        DEFAULT_FONT = _font
        break
else:
    DEFAULT_FONT = tkfont_det.nametofont("TkDefaultFont").actual()["family"]

import importlib.util
from openpyxl import load_workbook
import logging
import traceback

import platform
import subprocess

# Platform detection
SYSTEM = platform.system()
IS_WINDOWS = SYSTEM == "Windows"
IS_MACOS = SYSTEM == "Darwin"
IS_LINUX = SYSTEM == "Linux"

# python-docx import kontrolü
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("UYARI: python-docx kurulu değil! Kurulum için: pip install python-docx")


# win32com import kontrolü (sadece Windows için)
if IS_WINDOWS:
    try:
        import win32com.client
        WIN32_AVAILABLE = True
    except ImportError:
        WIN32_AVAILABLE = False
        print("UYARI: pywin32 kurulu değil! PDF dönüştürme özellikleri çalışmayacak.")
        print("Kurulum için: pip install pywin32")
else:
    WIN32_AVAILABLE = False
    # macOS/Linux için LibreOffice veya soffice kontrolü
    import shutil
    _libre = shutil.which("libreoffice") or shutil.which("soffice")
    # macOS default bundle path
    if not _libre and IS_MACOS:
        default_soffice = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if os.path.exists(default_soffice):
            _libre = default_soffice
    if _libre:
        LIBREOFFICE_BINARY = _libre
        try:
            result = subprocess.run([_libre, "--version"], capture_output=True, text=True, timeout=5)
            LIBREOFFICE_AVAILABLE = (result.returncode == 0)
            if LIBREOFFICE_AVAILABLE:
                print(f"✅ LibreOffice tespit edildi ({_libre}) - PDF dönüştürme aktif")
            else:
                print("UYARI: LibreOffice bulunamadı! PDF dönüştürme için LibreOffice kurmanız gerekiyor.")
        except (subprocess.TimeoutExpired, FileNotFoundError):
            LIBREOFFICE_AVAILABLE = False
            print("UYARI: LibreOffice bulunamadı! PDF dönüştürme için LibreOffice kurmanız gerekiyor.")
    else:
        LIBREOFFICE_AVAILABLE = False
        LIBREOFFICE_BINARY = None
        print("UYARI: LibreOffice bulunamadı! PDF dönüştürme için LibreOffice kurmanız gerekiyor.")

# Loglama ayarları
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('evrak_generator.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)


class DocumentProcessor:
    """Belge işleme sınıfı"""
    
    @staticmethod
    def sanitize_filename(name):
        """Dosya adını güvenli hale getirir - Unicode karakterleri destekler"""
        # Unicode normalize et
        name = unicodedata.normalize('NFKC', name)
        # Tehlikeli karakterleri kaldır
        sanitized = re.sub(r'[\\/:\*\?"<>\|]', '', name.strip())
        # Boş string kontrolü
        return sanitized if sanitized else "UNNAMED"
    
    @staticmethod
    def normalize_text_for_comparison(text):
        """Türkçe karakterleri ASCII eşdeğerine dönüştürür karşılaştırmalar için"""
        text = unicodedata.normalize('NFKD', text)
        # Türkçe ve Unicode karakterlerin ASCII eşdeğerleri
        replacements = {
            'Ç': 'C', 'ç': 'c',
            'Ğ': 'G', 'ğ': 'g', 
            'İ': 'I', 'ı': 'i', 'ì': 'i', 'í': 'i', 'î': 'i', 'ï': 'i',
            'Ö': 'O', 'ö': 'o', 'ò': 'o', 'ó': 'o', 'ô': 'o', 'õ': 'o',
            'Ş': 'S', 'ş': 's',
            'Ü': 'U', 'ü': 'u', 'ù': 'u', 'ú': 'u', 'û': 'u',
            # Arapça karakterler için temel eşleştirmeler
            'ا': 'a', 'ب': 'b', 'ت': 't', 'ث': 'th', 'ج': 'j', 'ح': 'h',
            'خ': 'kh', 'د': 'd', 'ذ': 'dh', 'ر': 'r', 'ز': 'z', 'س': 's',
            'ش': 'sh', 'ص': 's', 'ض': 'd', 'ط': 't', 'ظ': 'z', 'ع': 'a',
            'غ': 'gh', 'ف': 'f', 'ق': 'q', 'ك': 'k', 'ل': 'l', 'م': 'm',
            'ن': 'n', 'ه': 'h', 'و': 'w', 'ي': 'y'
        }
        
        # Unicode diacritikleri kaldır ve özel karakterleri değiştir
        result = ''.join(
            replacements.get(c, c) if not unicodedata.combining(c) else ''
            for c in text
        )
        return result.upper()
    
    @staticmethod
    def safe_string_comparison(text1, text2):
        """Unicode-safe string karşılaştırma"""
        if not text1 or not text2:
            return False
        # Hem orijinal hem de normalize edilmiş versiyonları kontrol et
        norm1 = DocumentProcessor.normalize_text_for_comparison(text1)
        norm2 = DocumentProcessor.normalize_text_for_comparison(text2)
        return (text1.upper() == text2.upper()) or (norm1 == norm2) or (text2.upper() in text1.upper()) or (norm2 in norm1)
    
    @staticmethod
    def process_word_document(src_path, dst_path, replacements):
        """Word belgesini python-docx ile işler"""
        if not DOCX_AVAILABLE:
            logging.error("python-docx kurulu değil!")
            return False
        
        logging.info(f"Word işleme başladı: {os.path.basename(src_path)}")
        
        try:
            # Belgeyi aç
            doc = Document(src_path)
            replacement_count = 0
            
            # Tüm metin içeriklerini işle
            text_elements = []
            
            # Paragrafları topla
            for paragraph in doc.paragraphs:
                text_elements.append(('paragraph', paragraph))
            
            # Tabloları topla
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text_elements.append(('table', paragraph))
            
            # Header'ları topla
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text_elements.append(('header', paragraph))
                        
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text_elements.append(('footer', paragraph))
            
            # Tüm elementleri işle
            for element_type, paragraph in text_elements:
                original_text = paragraph.text
                new_text = original_text
                
                # Değiştirmeleri uygula
                for key, value in replacements.items():
                    if key and key in new_text:
                        new_text = new_text.replace(key, str(value) if value else "")
                        replacement_count += 1
                        logging.debug(f"{element_type} - {key} -> {value}")
                
                # Değişiklik varsa güncelle
                if new_text != original_text:
                    # Paragrafın tüm run'larını temizle ve yeni metni yaz
                    for run in paragraph.runs:
                        run.text = ""
                    
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)
            
            logging.info(f"Toplam {replacement_count} değişiklik yapıldı")
            
            # Belgeyi kaydet
            doc.save(dst_path)
            logging.info(f"Belge kaydedildi: {dst_path}")
            
            return True
            
        except Exception as e:
            logging.error(f"Word işleme hatası: {e}")
            logging.error(traceback.format_exc())
            return False
    
    @staticmethod
    def process_excel_document(src_path, dst_path, replacements):
        """Excel belgesini işler"""
        logging.info(f"Excel işleme başladı: {os.path.basename(src_path)}")
        
        try:
            wb = load_workbook(src_path)
            replacement_count = 0
            
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str):
                            original_text = cell.value
                            new_text = original_text
                            
                            # Önce replacement dictionary'deki anahtarları değiştir
                            for key, value in replacements.items():
                                if key in new_text:
                                    new_text = new_text.replace(key, str(value))
                                    replacement_count += 1
                            
                            # Sonra replacement'ta olmayan placeholder'ları tamamen sil
                            # Faaliyet tarihi placeholder'ını özel olarak kontrol et
                            if "[DEĞİŞTİR:FAALİYETTARİH]" in new_text and "[DEĞİŞTİR:FAALİYETTARİH]" not in replacements:
                                new_text = new_text.replace("[DEĞİŞTİR:FAALİYETTARİH]", "")
                                replacement_count += 1
                                logging.info("Faaliyet tarihi placeholder'ı Excel'den silindi")
                            
                            if new_text != original_text:
                                cell.value = new_text
            
            # Özel işleme: Yıllık Değerlendirme Raporu için RD yöntemi güncellemesi
            filename = os.path.basename(src_path)
            logging.info(f"Excel dosya adı kontrol ediliyor: '{filename}'")
            
            # Debug için tüm kontrolleri yaz - case insensitive
            filename_normalized = unicodedata.normalize('NFKC', filename.upper())
            check1 = DocumentProcessor.safe_string_comparison(filename_normalized, "YILLIK")
            check2 = (DocumentProcessor.safe_string_comparison(filename_normalized, "DEĞERLENDIRME") or 
                     DocumentProcessor.safe_string_comparison(filename_normalized, "DEGERLENDIRME"))
            check3 = DocumentProcessor.safe_string_comparison(filename_normalized, "RAPORU")
            logging.info(f"[DEBUG] Excel kontrolleri: Yıllık={check1}, Değerlendirme={check2}, Raporu={check3}")
            logging.info(f"[DEBUG] Normalized filename: '{filename_normalized}'")
            
            # Daha geniş kontrolle Yıllık Değerlendirme Raporu'nu yakala
            if check1 and check2 and check3:
                logging.info("Yıllık Değerlendirme Raporu tespit edildi - RD yöntemi güncelleniyor")
                DocumentProcessor.update_rd_method_in_excel(wb, replacements)
            else:
                logging.info("Normal Excel dosyası - RD güncelleme yok")
            
            wb.save(dst_path)
            logging.info(f"Excel kaydedildi. {replacement_count} değişiklik yapıldı.")
            return True
            
        except Exception as e:
            logging.error(f"Excel işleme hatası: {e}")
            return False
    
    @staticmethod
    def update_rd_method_in_excel(workbook, replacements):
        """Yıllık Değerlendirme Raporu'nda G15 hücresini günceller"""
        try:
            # RD yöntemini al
            rd_method = replacements.get("[DEĞİŞTİR:RDYONTEMI]", "Matris")
            logging.info(f"Excel'de RD yöntemi güncellemesi başladı: {rd_method}")
            
            # Tüm worksheetlerde G15 hücresini kontrol et
            for sheet_index, ws in enumerate(workbook.worksheets):
                cell_g15 = ws['G15']
                cell_value = str(cell_g15.value) if cell_g15.value else ""
                
                logging.info(f"Sheet {sheet_index + 1} - G15 hücresi: '{cell_value}'")
                
                # METOD kelimesini farklı varyasyonlarda ara
                metod_variations = ["METOD", "metod", "Metod", "METHOD", "method", "Method"]
                
                for metod_var in metod_variations:
                    if metod_var in cell_value:
                        # METOD yazısını RD yöntemi ile değiştir
                        old_value = cell_value
                        new_value = cell_value.replace(metod_var, rd_method)
                        cell_g15.value = new_value
                        logging.info(f"G15 hücresi güncellendi: '{old_value}' -> '{new_value}' (RD Yöntemi: {rd_method})")
                        return
                
                # Eğer hiçbir şey bulunamazsa ve hücre boşsa, direkt RD yöntemini yaz
                if not cell_value or cell_value.strip() == "":
                    cell_g15.value = rd_method
                    logging.info(f"G15 hücresine RD yöntemi yazıldı (boş hücre): {rd_method}")
                    return
            
            # Hiçbir şey bulunamadıysa ilk sheet'in G15'ine yaz
            first_sheet = workbook.worksheets[0]
            first_sheet['G15'].value = rd_method
            logging.info(f"G15 hücresine RD yöntemi yazıldı (varsayılan): {rd_method}")
            
        except Exception as e:
            logging.error(f"RD yöntemi güncelleme hatası: {e}")
            import traceback
            logging.error(traceback.format_exc())


class PDFConverter:
    """PDF dönüştürme sınıfı"""
    
    @staticmethod
    def export_pdf_from_docx(docx_path, pdf_path):
        """
        Word belgesini PDF'e dönüştürür:
          - Windows + pywin32: mevcut COM yolu
          - Diğer platformlar: LibreOffice CLI headless
        """
        if IS_WINDOWS and WIN32_AVAILABLE:
            # COM kütüphanesini başlat (her thread için gerekli)
            import pythoncom
            pythoncom.CoInitialize()
            from win32com.client import DispatchEx
            word = DispatchEx('Word.Application')
            word.Visible = False
            word.DisplayAlerts = False

            doc = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=True)
            doc.ExportAsFixedFormat(
                OutputFileName=os.path.abspath(pdf_path),
                ExportFormat=17,  # wdExportFormatPDF
                OpenAfterExport=False,
                OptimizeFor=0,    # wdExportOptimizeForPrint
                Range=0           # wdExportAllDocument
            )
            doc.Close(False)
            word.Quit()
            # COM temizliğini yap
            pythoncom.CoUninitialize()
            logging.info(f"PDF oluşturuldu (Win32): {os.path.basename(pdf_path)}")
            return True
        else:
            # LibreOffice CLI ile dönüştürme (macOS/Linux)
            if not LIBREOFFICE_AVAILABLE:
                logging.error("LibreOffice bulunamadı! PDF dönüştürme yapılamıyor.")
                return False
            
            try:
                # Hedef dizini oluştur
                os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
                
                logging.info(f"LibreOffice komutu: {LIBREOFFICE_BINARY} --headless --convert-to pdf --outdir {os.path.dirname(pdf_path)} {docx_path}")
                subprocess.run([
                    LIBREOFFICE_BINARY,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", os.path.dirname(pdf_path),
                    os.path.abspath(docx_path)
                ], check=True, timeout=60)
                
                # LibreOffice genellikle dosya adını değiştirir, kontrol et
                expected_pdf = os.path.join(os.path.dirname(pdf_path), 
                                          os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
                if os.path.exists(expected_pdf) and expected_pdf != pdf_path:
                    shutil.move(expected_pdf, pdf_path)
                
                logging.info(f"PDF oluşturuldu (LibreOffice): {os.path.basename(pdf_path)}")
                return True
            except subprocess.TimeoutExpired:
                logging.error("PDF dönüştürme zaman aşımı")
                return False
            except Exception as e:
                logging.error(f"PDF dönüştürme hatası (LibreOffice): {e}")
                return False    
    @staticmethod
    def export_pdf_from_xlsx(xlsx_path, pdf_path):
        """
        Excel belgesini PDF'e dönüştürür:
          - Windows + pywin32: mevcut COM yolu
          - Diğer platformlar: LibreOffice CLI headless
        """
        if IS_WINDOWS and WIN32_AVAILABLE:
            # COM kütüphanesini başlat (her thread için gerekli)
            import pythoncom
            pythoncom.CoInitialize()
            from win32com.client import DispatchEx
            excel = DispatchEx('Excel.Application')
            excel.Visible = False
            excel.DisplayAlerts = False

            wb = excel.Workbooks.Open(os.path.abspath(xlsx_path), ReadOnly=True)
            wb.ExportAsFixedFormat(0, os.path.abspath(pdf_path))
            wb.Close(False)
            excel.Quit()
            # COM temizliğini yap
            pythoncom.CoUninitialize()
            logging.info(f"PDF oluşturuldu (Win32): {os.path.basename(pdf_path)}")
            return True
        else:
            # LibreOffice CLI ile dönüştürme (macOS/Linux)
            if not LIBREOFFICE_AVAILABLE:
                logging.error("LibreOffice bulunamadı! PDF dönüştürme yapılamıyor.")
                return False
            
            try:
                # Hedef dizini oluştur
                os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
                
                logging.info(f"LibreOffice komutu: {LIBREOFFICE_BINARY} --headless --convert-to pdf --outdir {os.path.dirname(pdf_path)} {xlsx_path}")
                subprocess.run([
                    LIBREOFFICE_BINARY,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", os.path.dirname(pdf_path),
                    os.path.abspath(xlsx_path)
                ], check=True, timeout=60)
                
                # LibreOffice genellikle dosya adını değiştirir, kontrol et
                expected_pdf = os.path.join(os.path.dirname(pdf_path), 
                                          os.path.splitext(os.path.basename(xlsx_path))[0] + ".pdf")
                if os.path.exists(expected_pdf) and expected_pdf != pdf_path:
                    shutil.move(expected_pdf, pdf_path)
                
                logging.info(f"PDF oluşturuldu (LibreOffice): {os.path.basename(pdf_path)}")
                return True
            except subprocess.TimeoutExpired:
                logging.error("PDF dönüştürme zaman aşımı")
                return False
            except Exception as e:
                logging.error(f"PDF dönüştürme hatası (LibreOffice): {e}")
                return False

class EvrakGenerator:
    """Ana evrak oluşturma sınıfı"""
    
    def __init__(self):
        self.processor = DocumentProcessor()
        self.pdf_converter = PDFConverter()
    
    def find_template_file(self, template_filename):
        """Platform uyumlu template dosya arama"""
        # Birincil yol: Evraklar/YILLIKLAR/dosya.xlsx
        template_path = os.path.normpath(os.path.join("Evraklar", "YILLIKLAR", template_filename))
        if os.path.exists(template_path):
            return template_path
        
        # İkincil yol: Çalışma dizini + Evraklar/YILLIKLAR/dosya.xlsx
        alt_path = os.path.normpath(os.path.join(os.getcwd(), "Evraklar", "YILLIKLAR", template_filename))
        if os.path.exists(alt_path):
            return alt_path
        
        # Üçüncül yol: YILLIKLAR klasöründeki dosyaları manuel arama (Unicode uyumlu)
        yilliklar_dir = os.path.normpath(os.path.join("Evraklar", "YILLIKLAR"))
        if os.path.exists(yilliklar_dir):
            try:
                files = os.listdir(yilliklar_dir)
                for file in files:
                    if DocumentProcessor.safe_string_comparison(file, template_filename):
                        return os.path.normpath(os.path.join(yilliklar_dir, file))
            except Exception as e:
                logging.error(f"YILLIKLAR klasör tarama hatası: {e}")
        
        # Alternatif yol için de manuel arama
        alt_yilliklar_dir = os.path.normpath(os.path.join(os.getcwd(), "Evraklar", "YILLIKLAR"))
        if os.path.exists(alt_yilliklar_dir):
            try:
                files = os.listdir(alt_yilliklar_dir)
                for file in files:
                    if DocumentProcessor.safe_string_comparison(file, template_filename):
                        return os.path.normpath(os.path.join(alt_yilliklar_dir, file))
            except Exception as e:
                logging.error(f"Alternatif YILLIKLAR klasör tarama hatası: {e}")
        
        logging.error(f"Template dosyası hiçbir yolda bulunamadı: {template_filename}")
        logging.error(f"Aranan yollar: {template_path}, {alt_path}")
        logging.error(f"Çalışma dizini: {os.getcwd()}")
        return None
    
    def load_replacements(self):
        """veri.xlsx dosyasından değiştirme verilerini yükler"""
        try:
            df = pd.read_excel("veri.xlsx", dtype=str, engine='openpyxl')
            logging.info("veri.xlsx başarıyla yüklendi")
            
            replacements = {}
            for _, row in df.iterrows():
                key = row["Anahtar"]
                value = row["Karşılık"]
                if pd.notna(key) and pd.notna(value):
                    replacements[str(key)] = str(value)
                elif pd.notna(key):
                    replacements[str(key)] = ""
            
            logging.info(f"Toplam {len(replacements)} anahtar yüklendi")
            return replacements, df
            
        except FileNotFoundError:
            logging.error("veri.xlsx bulunamadı")
            raise
        except Exception as e:
            logging.error(f"Veri yükleme hatası: {e}")
            raise
    
    def get_project_name(self, replacements):
        """
        Proje adını alır:
          • Proje Adı varsa onu kullanır,
          • Proje Adı yoksa Şirket Unvanı’nı kullanır,
          • İkisi de yoksa sabit olarak 'PROJE' döner.
        """
        proje_key   = "[DEĞİŞTİR:PROJEADI]"
        company_key = "[DEĞİŞTİR:ŞİRKET UNVANI]"   

        # 1) Önce Proje Adı’nı al
        proje = replacements.get(proje_key, "").strip()
        if not proje:
            # 2) Proje Adı boşsa Şirket Unvanı’nı al
            sirket = replacements.get(company_key, "").strip()
            # 3) Şirket Unvanı da boşsa 'PROJE'
            proje = sirket if sirket else "PROJE"

        # 4) Güvenli klasör/folder adı üret ve döndür
        return self.processor.sanitize_filename(proje)    
    def create_folders(self, project_name):
        """Gerekli klasörleri oluşturur - Unicode karakterleri destekler"""
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        
        # Unicode karakterleri destekleyen güvenli klasör adı
        safe_project_name = DocumentProcessor.sanitize_filename(project_name)
        target_folder = os.path.normpath(os.path.join(desktop, safe_project_name))
        
        today = datetime.datetime.now().strftime("%Y-%m-%d")
        backup_folder = os.path.normpath(os.path.join("yedekler", f"{today} - {safe_project_name}"))
        
        try:
            os.makedirs(target_folder, exist_ok=True)
            os.makedirs(backup_folder, exist_ok=True)
            logging.info(f"Klasörler oluşturuldu: {target_folder}, {backup_folder}")
        except Exception as e:
            logging.error(f"Klasör oluşturma hatası: {e}")
            # Windows encoding sorunları için alternatif yol
            if IS_WINDOWS:
                try:
                    target_folder = os.path.normpath(os.path.join(desktop, "PROJECT_OUTPUT"))
                    backup_folder = os.path.normpath(os.path.join("yedekler", f"{today} - PROJECT_BACKUP"))
                    os.makedirs(target_folder, exist_ok=True)
                    os.makedirs(backup_folder, exist_ok=True)
                    logging.info(f"Alternatif klasörler oluşturuldu: {target_folder}, {backup_folder}")
                except Exception as e2:
                    logging.error(f"Alternatif klasör oluşturma da başarısız: {e2}")
                    raise
            else:
                raise
        
        return target_folder, backup_folder
    
    def get_available_documents(self, rd_method=None):
        """İşlenebilir belgeleri listeler (RD yöntemine göre filtreleyebilir)"""
        evraklar_path = "Evraklar"
        if not os.path.exists(evraklar_path):
            os.makedirs(evraklar_path)
            logging.warning("Evraklar klasörü oluşturuldu")
            return []
        
        documents = []
        filtered_count = 0
        
        for filename in os.listdir(evraklar_path):
            if filename.lower().endswith((".docx", ".xlsx")) and not filename.startswith("~$"):
                # RD yöntemi filtresi
                should_include = True
                
                if rd_method:
                    logging.info(f"[DEBUG] Filtreleme kontrol: RD='{rd_method}' (len={len(rd_method)}), Dosya={filename}")
                    filename_upper = filename.upper().strip()
                    # Unicode normalization to handle Turkish characters
                    filename_normalized = unicodedata.normalize('NFKC', filename_upper)
                    
                    if rd_method == "Matris":
                        logging.info(f"[DEBUG] Matris dalında - filename_normalized='{filename_normalized}'")
                        # Matris seçiliyse Fine Kinney dosyasını hariç tut
                        fine_kinney_patterns = ["FINE_KINNEY", "FİNE_KINNEY", "FINE KINNEY", "FİNE KINNEY"]
                        contains_fine_kinney = any(pattern in filename_normalized for pattern in fine_kinney_patterns)
                        logging.info(f"[DEBUG] Fine Kinney kontrol: {contains_fine_kinney} - patterns: {fine_kinney_patterns}")
                        if contains_fine_kinney:
                            logging.info(f"[FILTER] Fine Kinney dosyası atlandı (Matris seçili): {filename}")
                            should_include = False
                            filtered_count += 1
                    elif rd_method == "Fine Kinney":
                        # Fine Kinney seçiliyse Matris dosyasını hariç tut
                        matris_patterns = ["MATRIS", "MATRİS"]
                        contains_matris = any(pattern in filename_normalized for pattern in matris_patterns)
                        contains_risk = "RİSK" in filename_normalized or "RISK" in filename_normalized
                        logging.info(f"[DEBUG] Matris kontrol: matris={contains_matris}, risk={contains_risk}")
                        if contains_matris and contains_risk:
                            logging.info(f"[FILTER] Matris dosyası atlandı (Fine Kinney seçili): {filename}")
                            should_include = False
                            filtered_count += 1
                
                if should_include:
                    documents.append(filename)
                    logging.debug(f"[INCLUDE] Dahil edildi: {filename}")
        
        # Yıllık plan seçeneklerini ekle
        documents = self.process_yearly_plan_options(documents)
        
        documents.sort()
        if rd_method:
            logging.info(f"{len(documents)} belge bulundu, {filtered_count} dosya filtrelendi (RD yöntemi: {rd_method})")
        else:
            logging.info(f"{len(documents)} belge bulundu (Filtresiz)")
        
        return documents
    
    def process_yearly_plan_options(self, documents):
        """Yıllık plan seçeneklerini işler"""
        processed_documents = []
        yearly_plans_found = []
        
        # YILLIKLAR klasöründeki dosyaları kontrol et
        yilliklar_path = os.path.join("Evraklar", "YILLIKLAR")
        logging.info(f"YILLIKLAR klasörü kontrol ediliyor: {yilliklar_path}")
        logging.info(f"Klasör var mı: {os.path.exists(yilliklar_path)}")
        
        if os.path.exists(yilliklar_path):
            try:
                files = os.listdir(yilliklar_path)
                logging.info(f"YILLIKLAR klasöründeki dosyalar: {files}")
                
                for filename in files:
                    if filename.lower().endswith(".xlsx") and not filename.startswith("~$"):
                        filename_upper = filename.upper()
                        logging.info(f"YILLIKLAR dosya kontrolü: {filename}")
                        # ASCII tabanlı kontrol için aksanları kaldır
                        filename_stripped = DocumentProcessor.normalize_text_for_comparison(filename)
                        # Plan türlerini tespit et
                        if DocumentProcessor.safe_string_comparison(filename_stripped, "EGITIM"):
                            yearly_plans_found.append("Yıllık Eğitim Planı")
                            logging.info("Yıllık Eğitim Planı bulundu!")
                        elif DocumentProcessor.safe_string_comparison(filename_stripped, "CALISMA"):
                            yearly_plans_found.append("Yıllık Çalışma Planı")
                            logging.info("Yıllık Çalışma Planı bulundu!")
                        elif (DocumentProcessor.safe_string_comparison(filename_stripped, "DEGERLENDIRME") and 
                              DocumentProcessor.safe_string_comparison(filename_stripped, "RAPORU")):
                            yearly_plans_found.append("Yıllık Değerlendirme Raporu")
                            logging.info("Yıllık Değerlendirme Raporu bulundu!")
            except Exception as e:
                logging.error(f"YILLIKLAR klasörü okuma hatası: {e}")
        
        # Normal belgeleri ekle
        for doc in documents:
            doc_upper = doc.upper()
            
            # Yıllık plan dosyalarını atla (bunlar zaten YILLIKLAR klasöründe)
            if "YILLIK" in doc_upper and ("EĞİTİM" in doc_upper or "ÇALIŞMA" in doc_upper):
                continue
            else:
                processed_documents.append(doc)
        
        # Yıllık planları sadece bir kez ekle (set kullanarak)
        unique_plans = set(yearly_plans_found)
        if "Yıllık Eğitim Planı" in unique_plans:
            processed_documents.append("Yıllık Eğitim Planı")
        if "Yıllık Çalışma Planı" in unique_plans:
            processed_documents.append("Yıllık Çalışma Planı")
        if "Yıllık Değerlendirme Raporu" in unique_plans:
            processed_documents.append("Yıllık Değerlendirme Raporu")
        
        logging.info(f"Yıllık planlar bulundu: {yearly_plans_found}")
        return processed_documents
    
    def is_yearly_plan_document(self, filename):
        """Dosyanın yıllık plan belgesi olup olmadığını kontrol eder"""
        return filename in ["Yıllık Eğitim Planı", "Yıllık Çalışma Planı"]

    def is_yearly_report_document(self, filename):
        """Dosyanın yıllık değerlendirme raporu belgesi olup olmadığını kontrol eder"""
        return filename == "Yıllık Değerlendirme Raporu"
    
    def process_yearly_plan_document(self, filename, replacements, project_name, target_folder, backup_folder):
        """Yıllık plan belgesini özel algoritmayla işler"""
        logging.info(f"=== Yıllık plan belgesi işleniyor: {filename} ===")
        
        # Yıl kontrolü yap
        use_dynamic_algorithm = self.check_year_compatibility(replacements)
        logging.info(f"Dinamik algoritma kullanılacak mı: {use_dynamic_algorithm}")
        
        # Çalışan sayısını al
        calisanlar = self.get_calisanlar_sayisi(replacements)
        is_kurullu = calisanlar >= 50
        logging.info(f"Çalışan sayısı: {calisanlar}, Kurullu: {is_kurullu}")
        
        # Uygun template'i seç
        template_path = self.select_yearly_template(filename, is_kurullu)
        if not template_path:
            logging.error(f"Yıllık plan template bulunamadı: {filename}")
            return False
        
        logging.info(f"Seçilen template: {template_path}")
        
        # Plan türünü belirle
        if "EĞİTİM" in filename.upper() or "EĞITIM" in filename.upper():
            plan_type = "Eğitim Planı"
        else:
            plan_type = "Çalışma Planı"
        
        logging.info(f"Plan türü belirlendi: {plan_type} (dosya: {filename})")
        
        # Hedef dosya yolları
        kurullu_text = "KURULLU" if is_kurullu else "KURULSUZ"
        dst_filename = f"{project_name} - Yıllık {plan_type} {kurullu_text}.xlsx"
        dst_path = os.path.join(target_folder, dst_filename)
        backup_path = os.path.join(backup_folder, dst_filename)
        
        try:
            # Template'i kopyala
            shutil.copy2(template_path, dst_path)
            logging.info(f"Template kopyalandı: {template_path} -> {dst_path}")
            
            # Placeholder'ları doldur
            self.fill_excel_placeholders(dst_path, replacements)
            
            # Dinamik algoritma uygula (geçmiş ayları temizle)
            if use_dynamic_algorithm:
                logging.info("Dinamik algoritma uygulanıyor...")
                self.apply_dynamic_algorithm(dst_path, plan_type,
                                             replacements.get("[DEĞİŞTİR:YILLIK:TARİH]", None))
                logging.info("Dinamik algoritma tamamlandı.")
            # Yıllık silme kuralları her durumda uygula
            logging.info("Yıllık silme kuralları uygulanıyor...")
            self.apply_yearly_deletion_rules(dst_path, plan_type, replacements)
            logging.info("Yıllık silme kuralları tamamlandı.")
            
            # Yedek kopyala
            shutil.copy(dst_path, backup_path)
            
            # PDF oluşturma tercihi
            if getattr(self, 'generate_pdf', False):
                pdf_dir = os.path.join(target_folder, "PDF")
                os.makedirs(pdf_dir, exist_ok=True)
                pdf_filename = f"{project_name} - Yıllık {plan_type} {kurullu_text}.pdf"
                pdf_path = os.path.join(pdf_dir, pdf_filename)
                self.pdf_converter.export_pdf_from_xlsx(dst_path, pdf_path)
                logging.info(f"Yıllık plan PDF oluşturuldu: {pdf_path}")
            logging.info(f"Yıllık plan belgesi başarıyla oluşturuldu: {dst_filename}")
            return True
            
        except Exception as e:
            logging.error(f"Yıllık plan belgesi hatası: {e}")
            return False
    def process_yearly_report_document(self, filename, replacements, project_name, target_folder, backup_folder):
        """Yıllık değerlendirme raporu belgesini işler"""
        logging.info(f"=== Yıllık Değerlendirme Raporu işleniyor: {filename} ===")
        template_path = os.path.join("Evraklar", "YILLIKLAR", "YILLIK DEĞERLENDİRME RAPORU.xlsx")
        if not os.path.exists(template_path):
            logging.error(f"Yıllık Değerlendirme Raporu template bulunamadı: {template_path}")
            return False
        dst_filename = f"{project_name} - Yıllık Değerlendirme Raporu.xlsx"
        dst_path = os.path.join(target_folder, dst_filename)
        backup_path = os.path.join(backup_folder, dst_filename)
        try:
            shutil.copy2(template_path, dst_path)
            self.fill_excel_placeholders(dst_path, replacements)
            shutil.copy(dst_path, backup_path)
            pdf_dir = os.path.join(target_folder, "PDF")
            os.makedirs(pdf_dir, exist_ok=True)
            pdf_filename = f"{project_name} - Yıllık Değerlendirme Raporu.pdf"
            pdf_path = os.path.join(pdf_dir, pdf_filename)
            if getattr(self, 'generate_pdf', False):
                self.pdf_converter.export_pdf_from_xlsx(dst_path, pdf_path)
                logging.info(f"Yıllık Değerlendirme Raporu PDF oluşturuldu: {pdf_path}")
            logging.info(f"Yıllık Değerlendirme Raporu başarıyla oluşturuldu: {dst_filename}")
            return True
        except Exception as e:
            logging.error(f"Yıllık Değerlendirme Raporu hatası: {e}")
            return False

    def check_year_compatibility(self, replacements):
        """Yıl uyumluluğunu kontrol eder"""
        try:
            # YILLIK:YIL değerini al
            yillik_yil = replacements.get("[DEĞİŞTİR:YILLIK:YIL]", "")
            if not yillik_yil or yillik_yil == "nan":
                return False
            
            # YILLIK:TARİH ve YILLIK:YIL karşılaştırması
            tarih_str = replacements.get("[DEĞİŞTİR:YILLIK:TARİH]", "").strip()
            yil_str = replacements.get("[DEĞİŞTİR:YILLIK:YIL]", "").strip()
            if not tarih_str or not yil_str:
                logging.info("Yıl uyumluluğu kontrolü için TARİH veya YIL verisi eksik")
                return False
            try:
                tarih_obj = datetime.datetime.strptime(tarih_str, "%d.%m.%Y")
                tarih_yil = str(tarih_obj.year)
            except Exception:
                logging.error(f"Tarih parse edilemedi: {tarih_str}")
                return False
            result = yil_str == tarih_yil
            logging.info(f"Yıl kontrolü: YIL={yil_str} vs TARİH yılı={tarih_yil} -> {result}")
            return result
            
        except Exception as e:
            logging.error(f"Yıl kontrolü hatası: {e}")
            return False
    
    def get_calisanlar_sayisi(self, replacements):
        """Çalışan sayısını döndürür"""
        try:
            calisanlar_str = replacements.get("[DEĞİŞTİR:ÇALIŞANSAYISI]", "0")
            if not calisanlar_str or calisanlar_str == "nan":
                return 0
            return int(calisanlar_str)
        except:
            return 0
    
    def select_yearly_template(self, filename, is_kurullu):
        """Yıllık plan template'ini seçer"""
        logging.info(f"Template seçimi için dosya: {filename}")
        
        # ASCII tabanlı kontrol için aksanları kaldır ve normalize et
        filename_stripped = DocumentProcessor.normalize_text_for_comparison(filename)
        # Template seçimi
        if (DocumentProcessor.safe_string_comparison(filename_stripped, "DEGERLENDIRME") and 
            DocumentProcessor.safe_string_comparison(filename_stripped, "RAPORU")):
            logging.info("Değerlendirme raporu template'i seçiliyor")
            template_path = os.path.join("Evraklar", "YILLIKLAR", "YILLIK DEĞERLENDİRME RAPORU.xlsx")
        elif DocumentProcessor.safe_string_comparison(filename_stripped, "EGITIM"):
            logging.info("Eğitim planı template'i seçiliyor")
            if is_kurullu:
                template_path = os.path.join("Evraklar", "YILLIKLAR", "YILLIK EĞİTİM PLANI KURULLU.xlsx")
            else:
                template_path = os.path.join("Evraklar", "YILLIKLAR", "YILLIK EĞİTİM PLANI KURULSUZ.xlsx")
        elif DocumentProcessor.safe_string_comparison(filename_stripped, "CALISMA"):
            logging.info("Çalışma planı template'i seçiliyor")
            if is_kurullu:
                template_path = os.path.join("Evraklar", "YILLIKLAR", "YILLIK ÇALIŞMA PLANI KURULLU.xlsx")
            else:
                template_path = os.path.join("Evraklar", "YILLIKLAR", "YILLIK ÇALIŞMA PLANI KURULSUZ.xlsx")
        else:
            logging.error(f"Yıllık plan template tipi belirlenemedi: {filename}")
            return None
        
        # Platformlar arası uyumlu dosya yolu oluştur
        template_path = os.path.normpath(template_path)
        
        # Dosya varlığını kontrol et
        exists = os.path.exists(template_path)
        if not exists:
            # Windows'ta encoding sorunları için alternatif yolları dene
            alt_path = os.path.normpath(os.path.join(os.getcwd(), "Evraklar", "YILLIKLAR", os.path.basename(template_path)))
            exists = os.path.exists(alt_path)
            if exists:
                template_path = alt_path
        
        logging.info(f"Template seçimi: {template_path}")
        logging.info(f"Dosya var mı: {exists}")
        logging.info(f"Çalışma dizini: {os.getcwd()}")
        
        # YILLIKLAR klasörü içeriğini debug için listele
        yilliklar_dir = os.path.normpath(os.path.join("Evraklar", "YILLIKLAR"))
        if os.path.exists(yilliklar_dir):
            try:
                files = os.listdir(yilliklar_dir)
                logging.info(f"YILLIKLAR klasörü içeriği: {files}")
                # Aranan dosyanın varlığını manuel kontrol et
                template_basename = os.path.basename(template_path)
                for file in files:
                    if DocumentProcessor.safe_string_comparison(file, template_basename):
                        logging.info(f"Manuel kontrol: {file} eşleşti {template_basename} ile")
                        # Doğru dosya adıyla yolu yeniden oluştur
                        template_path = os.path.normpath(os.path.join(yilliklar_dir, file))
                        exists = True
                        break
            except Exception as e:
                logging.error(f"YILLIKLAR klasörü listeleme hatası: {e}")
        
        # Platform uyumlu template dosya arama kullan
        if 'template_path' not in locals():
            # Fonksiyon başından itibaren yeniden yazılacak
            template_filename = None
            if (DocumentProcessor.safe_string_comparison(filename_stripped, "DEGERLENDIRME") and 
                DocumentProcessor.safe_string_comparison(filename_stripped, "RAPORU")):
                template_filename = "YILLIK DEĞERLENDİRME RAPORU.xlsx"
            elif DocumentProcessor.safe_string_comparison(filename_stripped, "EGITIM"):
                if is_kurullu:
                    template_filename = "YILLIK EĞİTİM PLANI KURULLU.xlsx"
                else:
                    template_filename = "YILLIK EĞİTİM PLANI KURULSUZ.xlsx"
            elif DocumentProcessor.safe_string_comparison(filename_stripped, "CALISMA"):
                if is_kurullu:
                    template_filename = "YILLIK ÇALIŞMA PLANI KURULLU.xlsx"
                else:
                    template_filename = "YILLIK ÇALIŞMA PLANI KURULSUZ.xlsx"
            
            if template_filename:
                found_path = self.find_template_file(template_filename)
                return found_path
            else:
                return None
        
        # Eski yöntem fallback olarak - template_path varsa onu kullan
        return template_path if 'template_path' in locals() and exists else None
    
    def apply_dynamic_algorithm(self, excel_path, plan_type, tarih_str=None):
        """Dinamik algoritma - geçmiş ayları temizle (kullanıcı tarihine göre)"""
        try:
            from openpyxl import load_workbook
            from openpyxl.styles import PatternFill, Font

            # Başlık satırını tespit et
            header_row = 17 if "Eğitim" in plan_type else 6
            logging.info(f"Dinamik algoritma başlıyor - Plan türü: {plan_type}, Başlık satırı: {header_row}")

            # Kullanıcı tarihinden ayı al (fallback olarak bugün)
            try:
                if tarih_str:
                    dt = datetime.datetime.strptime(tarih_str, "%d.%m.%Y")
                else:
                    dt = datetime.datetime.now()
                current_month = dt.month
            except Exception:
                current_month = datetime.datetime.now().month
            logging.info(f"Dinamik tarih ayı: {current_month}")
            
            # Excel dosyasını aç
            wb = load_workbook(excel_path)
            ws = wb.active
            
            # Ay başlıklarını bul - tüm satırları tara
            ay_isimleri = ["OCAK", "ŞUBAT", "MART", "NİSAN", "MAYIS", "HAZİRAN", 
                          "TEMMUZ", "AĞUSTOS", "EYLÜL", "EKİM", "KASIM", "ARALIK"]
            
            # Önce doğru satırı bul
            found_header_row = None
            for row in range(1, 25):  # İlk 25 satırı kontrol et
                for col in range(1, ws.max_column + 1):
                    cell_value = ws.cell(row=row, column=col).value
                    if cell_value and str(cell_value).strip().upper() in ay_isimleri:
                        found_header_row = row
                        logging.info(f"Ay başlıkları {row}. satırda bulundu!")
                        break
                if found_header_row:
                    break
            
            if found_header_row:
                header_row = found_header_row
                logging.info(f"Güncellenen başlık satırı: {header_row}")
            
            cleaned_count = 0
            # Geçmiş ayları temizle
            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=header_row, column=col).value
                if cell_value and str(cell_value).strip().upper() in ay_isimleri:
                    ay_index = ay_isimleri.index(str(cell_value).strip().upper()) + 1
                    logging.info(f"Ay bulundu: {cell_value}, Index: {ay_index}, Kolon: {col}")
                    
                    # Eğer bu ay geçmiş aysa, sütunu temizle
                    if ay_index < current_month:
                        self.clear_column_content(ws, col, header_row)
                        cleaned_count += 1
                        logging.info(f"Geçmiş ay temizlendi: {cell_value}")
            
            logging.info(f"Toplam {cleaned_count} ay temizlendi")
            
            # Kaydet
            wb.save(excel_path)
            wb.close()
            
        except Exception as e:
            logging.error(f"Dinamik algoritma hatası: {e}")
            import traceback
            logging.error(traceback.format_exc())
    
    def clear_column_content(self, worksheet, column_index, header_row):
        """Sütun içeriğini temizler (başlık hariç)"""
        try:
            from openpyxl.styles import PatternFill, Font
            
            cleared_count = 0
            # Başlık satırından sonraki tüm hücreleri temizle
            for row in range(header_row + 1, worksheet.max_row + 1):
                cell = worksheet.cell(row=row, column=column_index)
                
                # X işaretlerini temizle (büyük/küçük harf, boşluk toleransı)
                if cell.value:
                    cell_value = str(cell.value).strip().upper()
                    if cell_value == "X" or cell_value == "x" or cell_value == "X ":
                        original_value = cell.value
                        cell.value = None
                        # Renk ve dolguyu temizle
                        cell.fill = PatternFill()
                        cell.font = Font()
                        cleared_count += 1
                        logging.info(f"Satır {row}, Sütun {column_index}: '{original_value}' temizlendi")
            
            logging.info(f"Sütun {column_index}'da {cleared_count} hücre temizlendi")
                    
        except Exception as e:
            logging.error(f"Sütun temizleme hatası: {e}")
            import traceback
            logging.error(traceback.format_exc())

    def apply_yearly_deletion_rules(self, excel_path, plan_type, replacements):
        """Yıllık silme kurallarına göre belirli hücreleri temizler."""
        try:
            # Tarih bilgisini al
            tarih_str = replacements.get("[DEĞİŞTİR:YILLIK:TARİH]", "")
            if not tarih_str:
                logging.info("YILLIK:TARİH değeri bulunamadı, silme kuralları uygulanmayacak")
                return
            # Tarih formatı dd.mm.yyyy
            try:
                tarih = datetime.datetime.strptime(tarih_str.strip(), "%d.%m.%Y")
            except Exception:
                logging.error(f"Tarih parse edilemedi: {tarih_str}")
                return
            ay = tarih.month
            # Plan tipi anahtarını oluştur
            is_kurullu = self.get_calisanlar_sayisi(replacements) >= 50
            if "Çalışma Planı" in plan_type:
                base = "calisma_plani"
            else:
                base = "egitim_plani"
            kur_text = "kurullu" if is_kurullu else "kurulsuz"
            plan_key = f"yillik_{base}_{kur_text}"
            logging.info(f"Yıllık silme kuralı hesaplandı: ay={ay}, plan_tipi={plan_key}")
            # Kuralları oku
            df_rules = pd.read_csv("YILLIK_SILME_KURALLARI.csv", encoding='utf-8')
            df_match = df_rules[(df_rules["ay"] == ay) & (df_rules["plan_tipi"] == plan_key)]
            if df_match.empty:
                logging.info(f"Silme kuralı bulunamadı: ay={ay}, plan_tipi={plan_key}")
                return
            cell_str = df_match.iloc[0]["hucreler"]
            if not cell_str or pd.isna(cell_str):
                logging.info(f"Silinecek hücre yok: ay={ay}, plan_tipi={plan_key}")
                return
            cell_list = [c.strip() for c in cell_str.split(";") if c.strip()]
            logging.info(f"Silinecek hücreler listesi ({len(cell_list)}): {cell_list}")
            wb = load_workbook(excel_path)
            ws = wb.active
            from openpyxl.styles import PatternFill, Font
            removed = 0
            for ref in cell_list:
                cell = ws[ref]
                if cell.value:
                    cell.value = None
                    cell.fill = PatternFill()
                    cell.font = Font()
                    removed += 1
                    logging.info(f"Hücre silindi: {ref}")
            wb.save(excel_path)
            wb.close()
            logging.info(f"Toplam {removed} hücre silindi (YILLIK_SILME_KURALLARI)")
        except Exception as e:
            logging.error(f"Yıllık silme kuralları hatası: {e}")
            logging.error(traceback.format_exc())
    
    def fill_excel_placeholders(self, excel_path, replacements):
        """Excel dosyasındaki placeholder'ları doldurur"""
        try:
            # Doğrudan excel dosyasını işle
            success = self.processor.process_excel_document(excel_path, excel_path, replacements)
            
            if success:
                logging.info(f"Placeholder'lar dolduruldu: {excel_path}")
            else:
                logging.error(f"Placeholder doldurma başarısız: {excel_path}")
                
        except Exception as e:
            logging.error(f"Placeholder doldurma hatası: {e}")
            import traceback
            logging.error(traceback.format_exc())
    
    def process_document(self, filename, replacements, project_name, target_folder, backup_folder):
        """Tek bir belgeyi işler"""
        logging.info(f"\n=== İşlem başlıyor: {filename} ===")
        
        # Yıllık değerlendirme raporu veya plan kontrolü
        if self.is_yearly_report_document(filename):
            return self.process_yearly_report_document(filename, replacements, project_name, target_folder, backup_folder)
        if self.is_yearly_plan_document(filename):
            return self.process_yearly_plan_document(filename, replacements, project_name, target_folder, backup_folder)
        
        src_path = os.path.join("Evraklar", filename)
        if not os.path.isfile(src_path):
            logging.error(f"Kaynak dosya bulunamadı: {src_path}")
            return False
        
        # Hedef dosya yolları
        dst_filename = f"{project_name} - {filename}"
        dst_path = os.path.join(target_folder, dst_filename)
        backup_path = os.path.join(backup_folder, dst_filename)
        
        # Dosya tipine göre işle
        success = False
        if filename.lower().endswith(".docx"):
            success = self.processor.process_word_document(src_path, dst_path, replacements)
        elif filename.lower().endswith(".xlsx"):
            success = self.processor.process_excel_document(src_path, dst_path, replacements)
        
        if success:
            # Yedek kopyala
            try:
                shutil.copy(dst_path, backup_path)
                logging.info("Yedek kopyalandı")
            except Exception as e:
                logging.error(f"Yedek kopyalama hatası: {e}")
            
            # PDF oluşturma seçeneği varsa
            if getattr(self, 'generate_pdf', False):
                pdf_dir = os.path.join(target_folder, "PDF")
                os.makedirs(pdf_dir, exist_ok=True)
                pdf_filename = f"{project_name} - {os.path.splitext(filename)[0]}.pdf"
                pdf_path = os.path.join(pdf_dir, pdf_filename)
                if filename.lower().endswith(".docx"):
                    self.pdf_converter.export_pdf_from_docx(dst_path, pdf_path)
                elif filename.lower().endswith(".xlsx"):
                    self.pdf_converter.export_pdf_from_xlsx(dst_path, pdf_path)
        
        logging.info(f"=== İşlem tamamlandı: {filename} ===\n")
        return success
    
    def create_all_documents(self):
        """Tüm belgeleri oluşturur"""
        try:
            replacements, df = self.load_replacements()
            project_name = self.get_project_name(replacements)
            target_folder, backup_folder = self.create_folders(project_name)
            
            # veri.xlsx'i yedekle
            df.to_excel(os.path.join(backup_folder, "veri.xlsx"), index=False, engine='openpyxl')
            
            # RD yöntemini al
            rd_method = replacements.get("[DEĞİŞTİR:RDYONTEMI]", "Matris")
            
            documents = self.get_available_documents(rd_method)
            if not documents:
                messagebox.showwarning("Uyarı", "İşlenecek belge bulunamadı!")
                return
            
            success_count = 0
            for doc in documents:
                if self.process_document(doc, replacements, project_name, target_folder, backup_folder):
                    success_count += 1
            
            messagebox.showinfo("Tamamlandı", 
                              f"İşlem tamamlandı!\n\n"
                              f"• {success_count}/{len(documents)} belge başarıyla işlendi\n"
                              f"• Belgeler masaüstünde '{project_name}' klasöründe")
            
        except Exception as e:
            messagebox.showerror("Hata", f"İşlem sırasında hata oluştu:\n{str(e)}")
            logging.error(f"Genel hata: {e}")
            logging.error(traceback.format_exc())
    
    def create_selected_documents(self, selected_files, replacements, project_name, 
                                target_folder, backup_folder):
        """Seçili belgeleri oluşturur"""
        success_count = 0
        for doc in selected_files:
            if self.process_document(doc, replacements, project_name, target_folder, backup_folder):
                success_count += 1
        
        messagebox.showinfo("Tamamlandı", 
                          f"İşlem tamamlandı!\n\n"
                          f"• {success_count}/{len(selected_files)} belge başarıyla işlendi")


class EvrakGeneratorGUI:
    """Kullanıcı arayüzü sınıfı"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("EVRAK GENERATOR")
        self.root.geometry("560x830")
        self.root.configure(bg="#e0e0e0")
        
        self.generator = EvrakGenerator()
        # PDF oluşturma seçeneği (GUI üzerinden işaretlenebilir)
        self.generate_pdf_var = tk.BooleanVar(value=False)
        
        self.create_ui()
    
    def create_ui(self):
        """Ana arayüzü oluşturur"""
        # Ana başlık frame
        title_frame = tk.Frame(self.root, bg="#e0e0e0")
        title_frame.pack(pady=(30, 20))
        
        # Büyük ve dikkat çekici başlık
        tk.Label(title_frame, text="📋 EVRAK GENERATOR", 
                font=(DEFAULT_FONT, 32, "bold"),
                bg="#e0e0e0", fg="#1a237e").pack()
        
        # Alt başlık
        tk.Label(title_frame, text="Dökümanları bir tıkla hazırlayın.", 
                font=(DEFAULT_FONT, 14),
                bg="#e0e0e0", fg="#1a237e").pack(pady=(5, 0))
        
        # Butonlar
        btn_frame = tk.Frame(self.root, bg="#e0e0e0")
        btn_frame.pack(pady=35)
        # PDF çıktısı seçeneği (isteğe bağlı)
        pdf_chk = tk.Checkbutton(btn_frame, text="Pdf",
                                 variable=self.generate_pdf_var,
                                 bg="#e0e0e0", fg="#1a237e",
                                 selectcolor="#e0e0e0",
                                 font=(DEFAULT_FONT, 10))
        pdf_chk.pack(pady=(0, 10))
        
        buttons = [
            ("Form Bilgilerini Doldur", self.launch_form, "#1a237e"),
            ("📄 Tüm Belgeleri Oluştur", self.create_all_documents, "#1a237e"),
            ("📑 Seçerek Belge Oluştur", self.create_selected_documents, "#1a237e"),
            ("📆 Toplu Yıllık Oluştur", self.launch_batch_yearly, "#1a237e"),
            ("📋 Toplu Faaliyet Formu Oluştur", self.launch_batch_faaliyet, "#2e7d32"),
            ("📜 Evrak Geçmişi", self.open_history, "#1a237e")
        ]
        
        for text, command, color in buttons:
            btn = self.create_styled_button(btn_frame, text, command, color)
            btn.pack(pady=8)
        
        # Evrak efekti (alt kısımda)
        self.create_document_effect()
        
        # Alt bilgi
        tk.Label(self.root, text="Created by Hüseyin İLHAN", 
                font=(DEFAULT_FONT, 10),
                bg="#e0e0e0", fg="#1a237e").pack(side="bottom", pady=10)
    
    def create_document_effect(self):
        """Estetik evrak efekti oluşturur"""
        # Arka plan canvas
        canvas = tk.Canvas(self.root, bg="#e0e0e0", highlightthickness=0, height=80)
        canvas.pack(fill="x", padx=20, pady=(20, 0))
        
        # Evrak ikonları - merkezi tasarım
        documents = ["📄", "📋", "📑", "📊", "📈", "📝", "📃"]
        
        # Orta kısımda dalgalı dizilim
        center_x = 350  # Canvas ortası
        for i, doc in enumerate(documents):
            x = center_x - 150 + i * 45
            y = 40 + 15 * ((-1) ** i)  # Dalgalı efekt
            canvas.create_text(x, y, text=doc, font=(DEFAULT_FONT, 16), fill="#b0b0b0")
        
        # Kenar süslemeler
        for i in range(3):
            x_left = 50 + i * 25
            x_right = 600 + i * 25
            y = 20 + i * 15
            canvas.create_text(x_left, y, text="📄", font=(DEFAULT_FONT, 12), fill="#d0d0d0")
            canvas.create_text(x_right, y, text="📋", font=(DEFAULT_FONT, 12), fill="#d0d0d0")
    
    def create_styled_button(self, parent, text, command, color):
        """3D görünümlü buton oluşturur"""
        btn = tk.Button(parent, text=text, command=command,
                       font=(DEFAULT_FONT, 12, "bold"),
                       bg="#f0f0f0", fg="#1a237e",
                       activebackground="#e0e0e0", activeforeground="#1a237e",
                       width=40, height=3, bd=3, pady=5,
                       relief="raised")
        
        # macOS için özel ayarlar
        if IS_MACOS:
            btn.configure(highlightbackground="#f0f0f0")
        
        return btn
    
    def launch_form(self):
        """Form uygulamasını başlatır"""
        # Form modülünü başlatmak için FORMMODUL.py kullan
        script_path = os.path.join(os.getcwd(), "FORMMODULU.py")
        if not os.path.isfile(script_path):
            messagebox.showerror("Hata", "Form scripti bulunamadı!")
            return
        
        try:
            # Modülü dinamik olarak yükle
            spec = importlib.util.spec_from_file_location("form_module", script_path)
            form_module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(form_module)
            
            # Yeni pencerede formu aç
            top = Toplevel(self.root)
            form_module.SadeFormApp(top)
            
        except Exception as e:
            messagebox.showerror("Hata", f"Form yüklenemedi:\n{str(e)}")
            logging.error(f"Form yükleme hatası: {e}")
    
    def launch_batch_yearly(self):
        """Toplu yıllık oluşturma penceresini gösterir"""
        batch_win = Toplevel(self.root)
        batch_win.title("Toplu Yıllık Oluştur")
        # Tarih seçimi
        frm = tk.Frame(batch_win)
        frm.pack(padx=10, pady=10)
        # YILLIK:YIL, YILLIK:TARİH ve YDR:YIL giriş alanları
        tk.Label(frm, text="YEP‑YÇP Yıl:", font=(DEFAULT_FONT, 10, "bold")).grid(row=0, column=0, sticky="w")
        self.batch_year_var = tk.StringVar(value=datetime.datetime.now().strftime("%Y"))
        tk.Entry(frm, textvariable=self.batch_year_var, width=6).grid(row=0, column=1, sticky="w")
        tk.Label(frm, text="Yıllık Tarih:", font=(DEFAULT_FONT, 10)).grid(row=0, column=2, sticky="w")
        self.batch_date_var = tk.StringVar(value=datetime.datetime.now().strftime("%d.%m.%Y"))
        tk.Entry(frm, textvariable=self.batch_date_var, width=12).grid(row=0, column=3, sticky="w")
        tk.Label(frm, text="YDR Yıl:", font=(DEFAULT_FONT, 10)).grid(row=0, column=4, sticky="w")
        self.batch_ydr_year_var = tk.StringVar(value=datetime.datetime.now().strftime("%Y"))
        tk.Entry(frm, textvariable=self.batch_ydr_year_var, width=6).grid(row=0, column=5, sticky="w")
        # Firma unvanlarını tek seferde getir
        tk.Button(frm, text="Seç Firmalar", command=self.select_companies,
                  font=(DEFAULT_FONT, 10, "bold"), width=10).grid(row=0, column=6, padx=5)
        # Tablo başlıkları
        headers = ["SGK No", "RD Yöntemi", "RD Tarih", "Telefon", "E-mail"]
        for c, h in enumerate(headers):
            tk.Label(frm, text=h, font=(DEFAULT_FONT, 10, "bold")).grid(row=1, column=c, padx=5, pady=5)
        # Satırlar için frame ve liste
        self.batch_rows_frame = frm
        self.batch_rows = []
        # Başlangıçta 10 satır ekle ve '+' butonunu hazırla
        for _ in range(10):
            self._add_batch_row()
        self.add_button = tk.Button(frm, text="+", command=self._add_batch_row,
                                    font=(DEFAULT_FONT, 12, "bold"), width=3)
        self.add_button.grid(row=2 + len(self.batch_rows)*2, column=0, pady=5)
        # Verileri kaydet (yıllıkverileri.xlsx)
        btn_save = self.create_styled_button(batch_win, "Verileri Kaydet", self.save_yearly_data, "#1a237e")
        btn_save.pack(pady=(5,2))
        btn_run = self.create_styled_button(batch_win, "Oluştur", self.run_batch_yearly, "#1a237e")
        btn_run.pack(pady=2)

    def _add_batch_row(self):
        """Yeni toplu yıllık satırı ekler"""
        if len(self.batch_rows) >= 20:
            messagebox.showwarning("Uyarı", "En fazla 20 firma girebilirsiniz!")
            return
        idx = len(self.batch_rows)
        row = idx * 2 + 2
        sgk_var = tk.StringVar()
        rd_method_var = tk.StringVar(value="Matris")
        rd_date_var = tk.StringVar(value=self.batch_date_var.get())
        phone_var = tk.StringVar()
        email_var = tk.StringVar()
        tk.Entry(self.batch_rows_frame, textvariable=sgk_var, width=12).grid(row=row, column=0, padx=5, pady=2)
        tk.OptionMenu(self.batch_rows_frame, rd_method_var, "Matris", "Fine Kinney").grid(row=row, column=1, padx=5, pady=2)
        tk.Entry(self.batch_rows_frame, textvariable=rd_date_var, width=12).grid(row=row, column=2, padx=5, pady=2)
        tk.Entry(self.batch_rows_frame, textvariable=phone_var, width=12).grid(row=row, column=3, padx=5, pady=2)
        tk.Entry(self.batch_rows_frame, textvariable=email_var, width=15).grid(row=row, column=4, padx=5, pady=2)
        # Firma unvanı göstermek için alt satır
        company_var = tk.StringVar()
        tk.Label(self.batch_rows_frame, textvariable=company_var, fg="#333").grid(
            row=row+1, column=0, columnspan=5, sticky="w", padx=5, pady=(0,4)
        )
        self.batch_rows.append((sgk_var, rd_method_var, rd_date_var, phone_var, email_var, company_var))

    def apply_company_info(self, sgk, replacements, tbl_path, nace_path):
        """FormModülü mantığıyla SGK No'ya göre firma verilerini ANKARA ve NACE tablolarından çeker."""
        try:
            # Dosya adı farklı olabilir; önce literal, sonra klasörde ANKARA içeren satırı bul
            if not os.path.exists(tbl_path):
                for f in os.listdir():
                    if "ANKARA" in f.upper() and f.lower().endswith(".xlsx"):
                        tbl_path = f
                        break
            df_ank = pd.read_excel(tbl_path, dtype=str, engine='openpyxl')
            sgk_col = "KISA SGK" if "KISA SGK" in df_ank.columns else df_ank.columns[15]
            df_ank[sgk_col] = df_ank[sgk_col].astype(str).str.strip()
            satir = df_ank[df_ank[sgk_col] == sgk]
            if not satir.empty:
                r = satir.iloc[0]
                field_map = {
                    "[DEĞİŞTİR:ŞİRKET UNVANI]": r.iloc[4],
                    "[DEĞİŞTİR:PROJEADI]": r.iloc[6],
                    "[DEĞİŞTİR:ADRES]": r.iloc[31],
                    "[DEĞİŞTİR:SGKSİCİL]": r.iloc[10],
                    "[DEĞİŞTİR:SGKSİCİL20PUNTO]": r.iloc[10],
                    "[DEĞİŞTİR:NACE]": r.iloc[9],
                    "[DEĞİŞTİR:TEHLİKESINIFI]": r.iloc[16],
                    "[DEĞİŞTİR:ÇALIŞANSAYISI]": r.iloc[19],
                    "[DEĞİŞTİR:ŞİRKET UNVANI20PUNTO]": r.iloc[4],
                    "[DEĞİŞTİR:PROJEADI20PUNTO]": r.iloc[6]
                }
                if "GRUP DIŞI" in str(r.iloc[4]).upper():
                    field_map["[DEĞİŞTİR:ŞİRKET UNVANI]"] = r.iloc[6]
                    field_map["[DEĞİŞTİR:ŞİRKET UNVANI20PUNTO]"] = r.iloc[6]
                    field_map["[DEĞİŞTİR:PROJEADI]"] = ""
                    field_map["[DEĞİŞTİR:PROJEADI20PUNTO]"] = ""
                field_map["[DEĞİŞTİR:ŞİRKETPROJE]"] = (
                    f"{field_map['[DEĞİŞTİR:ŞİRKET UNVANI]']} - "
                    f"{field_map.get('[DEĞİŞTİR:PROJEADI]', '')}"
                )
                replacements.update({k: str(v) for k, v in field_map.items()})
                # benzer şekilde NACE tablosunu bul
                if not os.path.exists(nace_path):
                    for f in os.listdir():
                        if "NACE" in f.upper() and f.lower().endswith(".xlsx"):
                            nace_path = f
                            break
                df_nace = pd.read_excel(nace_path, dtype=str, engine='openpyxl')
                kcol, acol = df_nace.columns[0], df_nace.columns[1]
                df_nace[kcol] = df_nace[kcol].astype(str).str.strip()
                found = df_nace[df_nace[kcol] == str(field_map['[DEĞİŞTİR:NACE]'])]
                desc = str(found.iloc[0][acol]) if not found.empty else ""
                replacements["[DEĞİŞTİR:NACEFAALİYET]"] = desc
                if desc:
                    replacements["[DEĞİŞTİR:NACEVEFAALİYET]"] = (
                        f"{field_map['[DEĞİŞTİR:NACE]']} - {desc}"
                    )
        except Exception as e:
            logging.error(f"Batch firma bilgisi yükleme hatası: {e}")

    def apply_dynamic_fields(self, replacements):
        """Tehlike sınıfına göre yıllık ve RD periyot/saat hesaplamalarını yapar"""
        try:
            tehlike = replacements.get("[DEĞİŞTİR:TEHLİKESINIFI]", "").strip()
            # Yıl ekleme haritası
            yil_map = {"AZ TEHLİKELİ": 6, "TEHLİKELİ": 4, "ÇOK TEHLİKELİ": 2}
            ek_yil = yil_map.get(tehlike, 0)
            fmt = "%d.%m.%Y"
            # Tarih başlangıç bitiş
            pairs = [("[DEĞİŞTİR:RDEKİPATAMAEĞİTİMHAZIRLANMA]", "[DEĞİŞTİR:RDGEÇERLİLİK]"),
                     ("[DEĞİŞTİR:ADEPEK3ATAMAEĞİTİM]", "[DEĞİŞTİR:ADEPGEÇERLİLİK]")]
            for inp, outp in pairs:
                start = replacements.get(inp, "").strip()
                if start:
                    try:
                        dt = datetime.datetime.strptime(start, fmt)
                        end = (dt + relativedelta(years=ek_yil)).strftime(fmt)
                        replacements[outp] = end
                    except Exception:
                        pass
            # Yıllık saat
            igu = "4 SAAT" if tehlike == "AZ TEHLİKELİ" else "8 SAAT"
            replacements["[DEĞİŞTİR:YILLIK:İGU:SAAT]"] = igu
            ih = "4 SAAT" if tehlike in ("AZ TEHLİKELİ", "TEHLİKELİ") else "8 SAAT"
            replacements["[DEĞİŞTİR:YILLIK:İH:SAAT]"] = ih
            # RD periyot ve muayene periyot
            per = {
                "AZ TEHLİKELİ": ("6 Yılda 1", "5 Yılda 1"),
                "TEHLİKELİ": ("4 Yılda 1", "3 Yılda 1"),
                "ÇOK TEHLİKELİ": ("2 Yılda 1", "Yılda 1")
            }.get(tehlike, ("", ""))
            replacements["[DEĞİŞTİR:YDR:RDPERİYOT]"] = per[0]
            replacements["[DEĞİŞTİR:YDR:MUAYENEPERİYOT]"] = per[1]
        except Exception as e:
            logging.error(f"Dinamik alan hesaplama hatası: {e}")

    def select_companies(self):
        """Tüm SGK kutucuklarından firma bilgilerini alıp unvanları gösterir"""
        for row in self.batch_rows:
            sgk_var, _, _, _, _, company_var = row
            kod = sgk_var.get().strip()
            if not kod:
                company_var.set("")
                continue
            # Geçici replacements dict
            repl = {}
            # Firma adını getirmek için apply_company_info'tan yararlan
            self.apply_company_info(kod, repl, "ANKARA İŞYERI TABLOSU.xlsx", "Nace Kod Listesi.xlsx")
            unvan = repl.get("[DEĞİŞTİR:ŞİRKET UNVANI]", "")
            company_var.set(unvan)

    def save_yearly_data(self):
        """Ekrandaki verileri yearlyverileri.xlsx dosyasına kaydeder"""
        """Ekrandaki verileri alıp yıllıkverileri.xlsx dosyasına yazar."""
        try:
            # Şablon veri.xlsx'ten anahtarları al
            df_base = pd.read_excel("veri.xlsx", dtype=str, engine='openpyxl')
            df_out = pd.DataFrame({"Anahtar": df_base["Anahtar"]})
            # Her satır için replacements oluştur ve karşılıkları kolonlara yaz
            for idx, row in enumerate(self.batch_rows, start=1):
                if idx > 20:
                    break
                sgk_var, rd_m_var, rd_d_var, phone_var, email_var, _ = row
                kod = sgk_var.get().strip()
                if not kod:
                    # boş bırak
                    df_out[f"Karşılık{idx}"] = ""
                    continue
                repl = {}
                # Firma verisi
                self.apply_company_info(kod, repl, "ANKARA İŞYERI TABLOSU.xlsx", "Nace Kod Listesi.xlsx")
                # Yıllık/RD/Telefon/E-mail bilgileri
                repl.update({
                    "[DEĞİŞTİR:YILLIK:TARİH]": self.batch_date_var.get().strip(),
                    "[DEĞİŞTİR:YILLIK:YIL]": self.batch_year_var.get().strip(),
                    "[DEĞİŞTİR:RDYONTEMI]": rd_m_var.get(),
                    "[DEĞİŞTİR:YDR:TARİH]": rd_d_var.get().strip(),
                    "[DEĞİŞTİR:YDR:YIL]": self.batch_ydr_year_var.get().strip(),
                    "[DEĞİŞTİR:RDEKİPATAMAEĞİTİMHAZIRLANMA]": rd_d_var.get().strip(),
                    "[DEĞİŞTİR:TELEFON]": phone_var.get().strip(),
                    "[DEĞİŞTİR:MAİL]": email_var.get().strip()
                })
                # Dinamik saat, periyot ve geçerlilik hesapla
                self.apply_dynamic_fields(repl)
                # Mapla karşılıkları
                df_out[f"Karşılık{idx}"] = df_base["Anahtar"].map(lambda k: repl.get(k, ""))
            # Kaydet
            # Yıllık verileri program klasöründe sakla
            base_dir = os.getcwd()
            out_path = os.path.join(base_dir, "yıllıkverileri.xlsx")
            df_out.to_excel(out_path, index=False, engine='openpyxl')
            messagebox.showinfo("Kaydedildi", f"Yıllık verileri kaydedildi:\n{out_path}")
            logging.info(f"Yıllık verileri dosyası oluşturuldu: {out_path}")
        except Exception as e:
            logging.error(f"Yıllık verileri kaydetme hatası: {e}")
            messagebox.showerror("Hata", f"Veri kaydetme hatası:\n{e}")

    def run_batch_yearly(self):
        """Girilen verilerle toplu yıllıkları oluşturur"""
        # PDF oluşturma tercihini al
        self.generator.generate_pdf = self.generate_pdf_var.get()
        date_str = self.batch_date_var.get().strip()
        try:
            dt = datetime.datetime.strptime(date_str, "%d.%m.%Y")
            year_str = str(dt.year)
        except Exception:
            messagebox.showerror("Hata", f"Tarih formatı yanlış: {date_str}")
            return
        # Hedef klasör
        # Hedef klasör: proje dizini altında 'Yıllıklar'
        base_dir = os.getcwd()
        folder_name = f"{datetime.datetime.now().strftime('%Y-%m-%d')} Yıllıklar"
        out_folder = os.path.join(base_dir, folder_name)
        os.makedirs(out_folder, exist_ok=True)
        # Şablon tablo dosyaları
        tbl_path = "ANKARA İŞYERI TABLOSU.xlsx"
        nace_path = "Nace Kod Listesi.xlsx"
        # Eğer kayıtlı yıllıkverileri.xlsx varsa, ondan oku
        yfile = os.path.join(os.getcwd(), "yıllıkverileri.xlsx")
        if os.path.exists(yfile):
            df_year = pd.read_excel(yfile, dtype=str, engine='openpyxl')
            for idx in range(1, 21):
                key = f"Karşılık{idx}"
                if key not in df_year.columns:
                    break
                replacements = dict(zip(df_year['Anahtar'], df_year[key].fillna('').astype(str)))
                sgk = replacements.get("[DEĞİŞTİR:SGKSİCİL]", "").strip()
                if not sgk:
                    continue
                self.generator.process_yearly_plan_document(
                    "Yıllık Eğitim Planı", replacements, sgk, out_folder, out_folder)
                self.generator.process_yearly_plan_document(
                    "Yıllık Çalışma Planı", replacements, sgk, out_folder, out_folder)
                self.generator.process_document(
                    "Yıllık Değerlendirme Raporu.xlsx", replacements, sgk, out_folder, out_folder)
            messagebox.showinfo("Tamam", "Toplu yıllık oluşturma tamamlandı.")
            return
        # Her satır için oluştur (GUI girişleri)
        for sgk_var, rd_method_var, rd_date_var, phone_var, email_var, _ in self.batch_rows:
            sgk = sgk_var.get().strip()
            if not sgk:
                continue
            rd_method = rd_method_var.get()
            rd_date = rd_date_var.get().strip()
            phone = phone_var.get().strip()
            email = email_var.get().strip()
            # Temel verileri yükle
            replacements, _ = self.generator.load_replacements()
            # Firma bilgilerini ANKARA tablosundan çek (FormModülü mantığı)
            try:
                df_ank = pd.read_excel(tbl_path, dtype=str, engine='openpyxl')
                sgk_col = "KISA SGK" if "KISA SGK" in df_ank.columns else df_ank.columns[15]
                df_ank[sgk_col] = df_ank[sgk_col].astype(str).str.strip()
                satir = df_ank[df_ank[sgk_col] == sgk]
                if not satir.empty:
                    r = satir.iloc[0]
                    # Firma bilgileri - ANKARA tablosundan
                    field_map = {
                        "[DEĞİŞTİR:İL]": r.iloc[3],
                        "[DEĞİŞTİR:ŞİRKET UNVANI]": r.iloc[4],
                        "[DEĞİŞTİR:PROJEADI]": r.iloc[6],
                        "[DEĞİŞTİR:ADRES]": r.iloc[31],
                        "[DEĞİŞTİR:SGKSİCİL]": r.iloc[10],
                        "[DEĞİŞTİR:SGKSİCİL20PUNTO]": r.iloc[10],
                        "[DEĞİŞTİR:NACE]": r.iloc[9],
                        "[DEĞİŞTİR:TEHLİKESINIFI]": r.iloc[16],
                        "[DEĞİŞTİR:ÇALIŞANSAYISI]": r.iloc[19],
                        # Uzman ve Hekim adları; V (22), Z (26) sütunlarından
                        "[DEĞİŞTİR:UZMANADI]": r.iloc[21],
                        "[DEĞİŞTİR:HEKİMADI]": r.iloc[25],
                        "[DEĞİŞTİR:ŞİRKET UNVANI20PUNTO]": r.iloc[4],
                        "[DEĞİŞTİR:PROJEADI20PUNTO]": r.iloc[6]
                    }
                    # Grup dışı kontrolü
                    if "GRUP DIŞI" in str(r.iloc[4]).upper():
                        # Şirket unvanı yerine proje adı
                        field_map["[DEĞİŞTİR:ŞİRKET UNVANI]"] = r.iloc[6]
                        field_map["[DEĞİŞTİR:ŞİRKET UNVANI20PUNTO]"] = r.iloc[6]
                        field_map["[DEĞİŞTİR:PROJEADI]"] = ""
                        field_map["[DEĞİŞTİR:PROJEADI20PUNTO]"] = ""
                    # Şirket-Proje kombinasyonu
                    field_map["[DEĞİŞTİR:ŞİRKETPROJE]"] = f"{field_map['[DEĞİŞTİR:ŞİRKET UNVANI]']} - {field_map.get('[DEĞİŞTİR:PROJEADI]', '')}"
                    replacements.update({k: str(v) for k, v in field_map.items()})
                    # NACE açıklaması ve kombinasyon
                    df_nace = pd.read_excel(nace_path, dtype=str, engine='openpyxl')
                    kcol, acol = df_nace.columns[0], df_nace.columns[1]
                    df_nace[kcol] = df_nace[kcol].astype(str).str.strip()
                    find = df_nace[df_nace[kcol] == str(field_map['[DEĞİŞTİR:NACE]'])]
                    aciklama = str(find.iloc[0][acol]) if not find.empty else ""
                    replacements["[DEĞİŞTİR:NACEFAALİYET]"] = aciklama
                    # NACE ve faaliyet kombinasyonu
                    if aciklama:
                        replacements["[DEĞİŞTİR:NACEVEFAALİYET]"] = f"{field_map['[DEĞİŞTİR:NACE]']} - {aciklama}"
            except Exception as e:
                logging.error(f"Batch firma bilgisi yükleme hatası: {e}")
            # Yıllık/RD/Telefon/E-mail bilgileri (firma SGKSİCİL ANKARA tablosundan çekildi)
            replacements.update({
                "[DEĞİŞTİR:YILLIK:TARİH]": date_str,
                "[DEĞİŞTİR:YILLIK:YIL]": year_str,
                "[DEĞİŞTİR:RDYONTEMI]": rd_method,
                "[DEĞİŞTİR:YDR:TARİH]": rd_date,
                "[DEĞİŞTİR:RDEKİPATAMAEĞİTİMHAZIRLANMA]": rd_date,
                "[DEĞİŞTİR:TELEFON]": phone,
                "[DEĞİŞTİR:MAİL]": email
            })
            # Dinamik periyot, saat, geçerlilik tarihleri hesapla
            self.apply_dynamic_fields(replacements)
            # Replacements Excel dosyası oluştur (veri.xlsx şablonuna benzer)
            try:
                df_base = pd.read_excel("veri.xlsx", dtype=str, engine='openpyxl')
                df_base.loc[df_base["Anahtar"].isin(replacements.keys()), "Karşılık"] = (
                    df_base["Anahtar"].map(replacements)
                )
                temp_xls = os.path.join(out_folder, f"veri_{sgk}.xlsx")
                df_base.to_excel(temp_xls, index=False, engine='openpyxl')
                logging.info(f"Replacements dosyası oluşturuldu: {temp_xls}")
            except Exception as e:
                logging.error(f"Replacements dosyası oluşturma hatası: {e}")
            # Yıllık Eğitim ve Çalışma Planı
            self.generator.process_yearly_plan_document(
                "Yıllık Eğitim Planı", replacements, sgk, out_folder, out_folder)
            self.generator.process_yearly_plan_document(
                "Yıllık Çalışma Planı", replacements, sgk, out_folder, out_folder)
            # Yıllık Değerlendirme Raporu
            self.generator.process_document(
                "Yıllık Değerlendirme Raporu.xlsx", replacements, sgk, out_folder, out_folder)
        messagebox.showinfo("Tamam", "Toplu yıllık oluşturma tamamlandı.")
    
    def launch_batch_faaliyet(self):
        """Toplu Faaliyet Formu oluşturma penceresini açar"""
        self.batch_faaliyet_window = Toplevel(self.root)
        self.batch_faaliyet_window.title("Toplu Faaliyet Formu Oluştur")
        self.batch_faaliyet_window.geometry("900x700")
        self.batch_faaliyet_window.configure(bg="#f8f9fa")
        self.batch_faaliyet_window.transient(self.root)
        self.batch_faaliyet_window.grab_set()
        
        # Pencere ortala
        self.batch_faaliyet_window.geometry("900x700+{}+{}".format(
            self.root.winfo_x() + 50, self.root.winfo_y() + 50))
        
        self.create_batch_faaliyet_ui()
    
    def create_batch_faaliyet_ui(self):
        """Toplu faaliyet formu arayüzünü oluşturur"""
        default_font = DEFAULT_FONT
        
        # Başlık
        title_label = tk.Label(self.batch_faaliyet_window, text="Toplu Faaliyet Formu Oluştur", 
                              font=(default_font, 18, "bold"),
                              bg="#f8f9fa", fg="#2c3e50")
        title_label.pack(pady=(15, 20))
        
        # Ana frame
        main_frame = tk.Frame(self.batch_faaliyet_window, bg="#f8f9fa")
        main_frame.pack(fill="both", expand=True, padx=20)
        
        # Faaliyet tarihi seçimi
        self.create_batch_faaliyet_tarihi_section(main_frame)
        
        # Scrollable SGK listesi
        self.create_batch_sgk_list_section(main_frame)
        
        # Alt butonlar
        self.create_batch_faaliyet_buttons(main_frame)
        
        # İlk 5 SGK kutusu ekle
        self.batch_sgk_entries = []
        self.batch_sgk_labels = []
        for i in range(5):
            self.add_batch_sgk_entry()
    
    def create_batch_faaliyet_tarihi_section(self, parent):
        """Faaliyet tarihi seçim bölümü"""
        default_font = DEFAULT_FONT
        
        date_frame = tk.LabelFrame(parent, text="Faaliyet Tarihi", 
                                  bg="#f8f9fa", fg="#2c3e50",
                                  font=(default_font, 12, "bold"),
                                  bd=2, relief="groove")
        date_frame.pack(fill="x", pady=(0, 20))
        
        date_inner = tk.Frame(date_frame, bg="#f8f9fa")
        date_inner.pack(padx=15, pady=10)
        
        # Basit Entry widget (silinebilir tarih için)
        tk.Label(date_inner, text="Tarih (DD.MM.YYYY):", 
                bg="#f8f9fa", fg="#2c3e50",
                font=(default_font, 10)).pack(side="left", padx=(0, 10))
        
        self.batch_faaliyet_tarihi = tk.Entry(date_inner, width=15, 
                                             font=(default_font, 10))
        
        # Bugünkü tarihi varsayılan olarak ayarla
        today = datetime.datetime.now().strftime("%d.%m.%Y")
        self.batch_faaliyet_tarihi.insert(0, today)
        
        self.batch_faaliyet_tarihi.pack(side="left", padx=(0, 10))
        
        # Açıklama label'ı
        info_label = tk.Label(date_inner, text="(Tarihsiz form için bu alanı boşaltın)", 
                             bg="#f8f9fa", fg="#666",
                             font=(default_font, 8, "italic"))
        info_label.pack(side="left")
    
    def create_batch_sgk_list_section(self, parent):
        """SGK listesi bölümü"""
        default_font = DEFAULT_FONT
        
        # SGK listesi frame
        sgk_frame = tk.LabelFrame(parent, text="SGK Kodları (7 Hane)", 
                                 bg="#f8f9fa", fg="#2c3e50",
                                 font=(default_font, 12, "bold"),
                                 bd=2, relief="groove")
        sgk_frame.pack(fill="both", expand=True, pady=(0, 20))
        
        # Scrollable canvas
        canvas = tk.Canvas(sgk_frame, bg="#f8f9fa", highlightthickness=0)
        scrollbar = tk.Scrollbar(sgk_frame, orient="vertical", command=canvas.yview)
        self.batch_scrollable_frame = tk.Frame(canvas, bg="#f8f9fa")
        
        scrollbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.create_window((0, 0), window=self.batch_scrollable_frame, anchor="nw")
        
        # Mouse wheel scrolling
        def on_mousewheel(event):
            if IS_MACOS:
                canvas.yview_scroll(int(-1 * event.delta), "units")
            else:
                canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        if IS_MACOS:
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        else:
            canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        # Frame boyut güncellemesi
        def configure_scroll_region(event=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        self.batch_scrollable_frame.bind("<Configure>", configure_scroll_region)
        
        # + Butonu frame (en altta sabit)
        self.batch_add_button_frame = tk.Frame(self.batch_scrollable_frame, bg="#f8f9fa")
        self.batch_add_button_frame.pack(fill="x", pady=(10, 0))
        
        # + Butonu
        self.batch_add_sgk_button = tk.Button(self.batch_add_button_frame, text="+ SGK Ekle",
                                             command=self.add_batch_sgk_entry,
                                             bg="#e3f2fd", fg="#1976d2",
                                             font=(default_font, 11, "bold"),
                                             relief="raised", bd=2)
        
        if IS_MACOS:
            self.batch_add_sgk_button.configure(highlightbackground="#e3f2fd")
        
        self.batch_add_sgk_button.pack(pady=5)
    
    def create_batch_faaliyet_buttons(self, parent):
        """Alt butonlar"""
        default_font = DEFAULT_FONT
        
        button_frame = tk.Frame(parent, bg="#f8f9fa")
        button_frame.pack(fill="x", pady=(0, 15))
        
        # SEÇ butonu
        sec_btn = tk.Button(button_frame, text="SEÇ - Şirket Proje Doğrula",
                           command=self.validate_batch_sgk_codes,
                           bg="#fff3e0", fg="#f57c00",
                           font=(default_font, 11, "bold"),
                           width=20, relief="raised", bd=3)
        
        if IS_MACOS:
            sec_btn.configure(highlightbackground="#fff3e0")
        
        sec_btn.pack(side="left", padx=(0, 15))
        
        # OLUŞTUR butonu
        olustur_btn = tk.Button(button_frame, text="OLUŞTUR",
                               command=self.create_batch_faaliyet_forms,
                               bg="#e8f5e8", fg="#2e7d32",
                               font=(default_font, 12, "bold"),
                               width=20, height=2, relief="raised", bd=3)
        
        if IS_MACOS:
            olustur_btn.configure(highlightbackground="#e8f5e8")
        
        olustur_btn.pack(side="right")
    
    def add_batch_sgk_entry(self):
        """Yeni SGK girişi ekler"""
        default_font = DEFAULT_FONT
        entry_index = len(self.batch_sgk_entries)
        
        # Yeni SGK frame'i
        sgk_entry_frame = tk.Frame(self.batch_scrollable_frame, bg="#f8f9fa")
        sgk_entry_frame.pack(fill="x", padx=10, pady=2, before=self.batch_add_button_frame)
        
        # SGK numarası
        tk.Label(sgk_entry_frame, text=f"{entry_index + 1}. SGK:", 
                bg="#f8f9fa", fg="#1a237e",
                font=(default_font, 10, "bold"), width=8).pack(side="left", padx=(0, 5))
        
        sgk_entry = tk.Entry(sgk_entry_frame, width=15, font=(default_font, 10))
        sgk_entry.pack(side="left", padx=(0, 10))
        
        # Şirket-Proje label
        label_text = tk.Label(sgk_entry_frame, text="", 
                             bg="#f8f9fa", fg="#666",
                             font=(default_font, 9), width=50, anchor="w")
        label_text.pack(side="left")
        
        self.batch_sgk_entries.append(sgk_entry)
        self.batch_sgk_labels.append(label_text)
        
        # Scroll bölgesini güncelle
        self.batch_scrollable_frame.update_idletasks()
    
    def validate_batch_sgk_codes(self):
        """SGK kodlarını doğrular ve şirket-proje bilgilerini gösterir"""
        # Ankara tablosunu yükle
        df_ankara = self.load_ankara_table_for_batch()
        if df_ankara is None:
            return
        
        for i, (sgk_entry, label) in enumerate(zip(self.batch_sgk_entries, self.batch_sgk_labels)):
            sgk_kod = sgk_entry.get().strip()
            
            if not sgk_kod:
                label.config(text="", fg="#666")
                continue
            
            if not self.validate_sgk_code_format(sgk_kod):
                label.config(text="❌ Geçersiz SGK (7 hane olmalı)", fg="#d32f2f")
                continue
            
            # SGK koduna göre satırı bul
            satir = self.find_sgk_row_in_table(df_ankara, sgk_kod)
            if satir is None:
                label.config(text="❌ SGK kodu bulunamadı", fg="#d32f2f")
                continue
            
            # Şirket ve proje bilgilerini göster
            sirket = str(satir.iloc[4]) if not pd.isna(satir.iloc[4]) else ""
            proje = str(satir.iloc[6]) if not pd.isna(satir.iloc[6]) else ""
            
            if "GRUP DIŞI" in sirket.upper():
                text = f"✅ {proje}"
            else:
                text = f"✅ {sirket} - {proje}" if proje else f"✅ {sirket}"
            
            label.config(text=text, fg="#2e7d32")
    
    def load_ankara_table_for_batch(self):
        """Ankara tablosunu yükler"""
        try:
            df = pd.read_excel("ANKARA İŞYERİ TABLOSU.xlsx", dtype=str, engine='openpyxl')
            logging.info("Ankara tablosu batch için yüklendi")
            return df
        except Exception as e:
            messagebox.showerror("Hata", f"Ankara tablosu açılamadı:\n{e}")
            logging.error(f"Ankara tablosu yükleme hatası: {e}")
            return None
    
    def validate_sgk_code_format(self, kod):
        """SGK kodunu doğrular"""
        return len(kod) == 7 and kod.isdigit()
    
    def find_sgk_row_in_table(self, df_ankara, kod):
        """SGK koduna göre satırı bulur"""
        # SGK sütununu bul
        sgk_col = "KISA SGK" if "KISA SGK" in df_ankara.columns else df_ankara.columns[15]
        df_ankara[sgk_col] = df_ankara[sgk_col].astype(str).str.strip()
        
        # Satırı bul
        satirlar = df_ankara[df_ankara[sgk_col] == kod]
        if satirlar.empty:
            return None
        
        return satirlar.iloc[0]
    
    def create_batch_faaliyet_forms(self):
        """Toplu faaliyet formlarını oluşturur"""
        try:
            # Tarihi al (basit Entry widget'ından)
            selected_date = self.batch_faaliyet_tarihi.get().strip()
            
            # Tarih formatını doğrula (eğer boş değilse)
            if selected_date:
                try:
                    datetime.datetime.strptime(selected_date, "%d.%m.%Y")
                except ValueError:
                    messagebox.showerror("Hata", "Geçersiz tarih formatı! DD.MM.YYYY formatında giriniz.")
                    return
            
            # Geçerli SGK kodlarını al
            valid_sgk_codes = []
            for sgk_entry in self.batch_sgk_entries:
                sgk_kod = sgk_entry.get().strip()
                if sgk_kod and len(sgk_kod) == 7 and sgk_kod.isdigit():
                    valid_sgk_codes.append(sgk_kod)
            
            if not valid_sgk_codes:
                messagebox.showwarning("Uyarı", "En az bir geçerli SGK kodu giriniz!")
                return
            
            # Masaüstünde klasör oluştur
            if selected_date:
                folder_name = f"{selected_date} - Faaliyet Formları"
            else:
                folder_name = "Faaliyet Formları"
            
            desktop_path = self.get_desktop_path()
            output_folder = os.path.join(desktop_path, folder_name)
            
            try:
                os.makedirs(output_folder, exist_ok=True)
                logging.info(f"Çıkış klasörü oluşturuldu: {output_folder}")
            except Exception as e:
                logging.error(f"Klasör oluşturma hatası: {e}")
                messagebox.showerror("Hata", f"Klasör oluşturulamadı: {str(e)}")
                return
            
            # Her SGK kodu için faaliyet formu oluştur
            created_files = []
            for sgk_kod in valid_sgk_codes:
                file_path = self.create_single_batch_faaliyet_form(sgk_kod, selected_date, output_folder)
                if file_path:
                    created_files.append(file_path)
            
            if created_files:
                messagebox.showinfo("Başarılı", 
                                  f"{len(created_files)} adet faaliyet formu oluşturuldu!\n\n"
                                  f"Klasör: {folder_name}\n"
                                  f"Masaüstünde dosyalar hazır.")
                self.batch_faaliyet_window.destroy()
            else:
                messagebox.showerror("Hata", "Hiçbir form oluşturulamadı!")
        
        except Exception as e:
            logging.error(f"Toplu faaliyet formu hatası: {e}")
            messagebox.showerror("Hata", f"Form oluşturma hatası: {str(e)}")
    
    def get_desktop_path(self):
        """Masaüstü yolunu döndürür"""
        if IS_WINDOWS:
            return os.path.join(os.path.expanduser("~"), "Desktop")
        elif IS_MACOS:
            return os.path.join(os.path.expanduser("~"), "Desktop")
        else:
            return os.path.join(os.path.expanduser("~"), "Desktop")
    
    def create_single_batch_faaliyet_form(self, sgk_kod, faaliyet_tarihi, output_folder):
        """Tek bir SGK kodu için faaliyet formu oluşturur"""
        try:
            # Ankara tablosunu yükle
            df_ankara = self.load_ankara_table_for_batch()
            if df_ankara is None:
                return None
            
            # SGK koduna göre satırı bul
            satir = self.find_sgk_row_in_table(df_ankara, sgk_kod)
            if satir is None:
                logging.error(f"SGK kodu bulunamadı: {sgk_kod}")
                return None
            
            # Form verilerini oluştur
            replacements = self.create_replacements_from_row(satir, faaliyet_tarihi)
            
            # Faaliyet formu şablonunu kullan
            template_path = os.path.join("Evraklar", "FAALİYET FORMU.xlsx")
            if not os.path.exists(template_path):
                logging.error(f"Faaliyet formu şablonu bulunamadı: {template_path}")
                return None
            
            # Hedef dosya adını oluştur
            sirket_proje = replacements.get("[DEĞİŞTİR:ŞİRKETPROJE]", f"SGK-{sgk_kod}")
            output_name = f"{sirket_proje} - Faaliyet Formu.xlsx"
            output_path = os.path.join(output_folder, output_name)
            
            # Template'i kopyala
            shutil.copy2(template_path, output_path)
            
            # Placeholder'ları doldur
            success = DocumentProcessor.process_excel_document(template_path, output_path, replacements)
            if not success:
                logging.error(f"Placeholder doldurma başarısız: {output_path}")
                return None
            
            logging.info(f"Faaliyet formu oluşturuldu: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"Tek faaliyet formu hatası: {e}")
            return None
    
    def create_replacements_from_row(self, satir, faaliyet_tarihi):
        """Satır verilerinden replacement dictionary oluşturur"""
        replacements = {}
        
        # Temel bilgiler
        replacements["[DEĞİŞTİR:ŞİRKET UNVANI]"] = str(satir.iloc[4]) if not pd.isna(satir.iloc[4]) else ""
        replacements["[DEĞİŞTİR:PROJEADI]"] = str(satir.iloc[6]) if not pd.isna(satir.iloc[6]) else ""
        replacements["[DEĞİŞTİR:ADRES]"] = str(satir.iloc[31]) if not pd.isna(satir.iloc[31]) else ""
        replacements["[DEĞİŞTİR:SGKSİCİL]"] = str(satir.iloc[10]) if not pd.isna(satir.iloc[10]) else ""
        replacements["[DEĞİŞTİR:SGKSİCİL20PUNTO]"] = str(satir.iloc[10]) if not pd.isna(satir.iloc[10]) else ""
        replacements["[DEĞİŞTİR:NACE]"] = str(satir.iloc[9]) if not pd.isna(satir.iloc[9]) else ""
        replacements["[DEĞİŞTİR:TEHLİKESINIFI]"] = str(satir.iloc[16]) if not pd.isna(satir.iloc[16]) else ""
        replacements["[DEĞİŞTİR:ÇALIŞANSAYISI]"] = str(satir.iloc[19]) if not pd.isna(satir.iloc[19]) else ""
        replacements["[DEĞİŞTİR:ŞİRKET UNVANI20PUNTO]"] = str(satir.iloc[4]) if not pd.isna(satir.iloc[4]) else ""
        replacements["[DEĞİŞTİR:PROJEADI20PUNTO]"] = str(satir.iloc[6]) if not pd.isna(satir.iloc[6]) else ""
        
        # Yeni placeholder'lar
        replacements["[DEĞİŞTİR:İL]"] = str(satir.iloc[3]) if not pd.isna(satir.iloc[3]) else ""
        replacements["[DEĞİŞTİR:UZMANADI]"] = str(satir.iloc[21]) if not pd.isna(satir.iloc[21]) else ""
        replacements["[DEĞİŞTİR:HEKİMADI]"] = str(satir.iloc[25]) if not pd.isna(satir.iloc[25]) else ""
        
        # Faaliyet tarihi - tarih varsa ekle, yoksa hiç ekleme (placeholder silinsin)
        if faaliyet_tarihi:
            replacements["[DEĞİŞTİR:FAALİYETTARİH]"] = faaliyet_tarihi
        # Tarih yoksa placeholder'ı replacement'a hiç eklememiz yeterli, DocumentProcessor otomatik silecek
        
        # Grup dışı kontrolü
        sirket_unvani = replacements["[DEĞİŞTİR:ŞİRKET UNVANI]"]
        proje_adi = replacements["[DEĞİŞTİR:PROJEADI]"]
        
        if "GRUP DIŞI" in sirket_unvani.upper():
            replacements["[DEĞİŞTİR:ŞİRKET UNVANI]"] = proje_adi
            replacements["[DEĞİŞTİR:ŞİRKET UNVANI20PUNTO]"] = proje_adi
            replacements["[DEĞİŞTİR:PROJEADI]"] = ""
            replacements["[DEĞİŞTİR:PROJEADI20PUNTO]"] = ""
            replacements["[DEĞİŞTİR:ŞİRKETPROJE]"] = proje_adi
        else:
            if proje_adi:
                replacements["[DEĞİŞTİR:ŞİRKETPROJE]"] = f"{sirket_unvani} - {proje_adi}"
            else:
                replacements["[DEĞİŞTİR:ŞİRKETPROJE]"] = sirket_unvani
        
        # NACE açıklamasını doldur
        self.fill_nace_for_replacement(replacements)
        
        return replacements
    
    def fill_nace_for_replacement(self, replacements):
        """NACE açıklamasını replacement'a ekler"""
        nace_kod = replacements.get("[DEĞİŞTİR:NACE]", "")
        
        if not nace_kod:
            return
        
        try:
            # NACE tablosunu yükle
            df_nace = pd.read_excel("Nace Kod Listesi.xlsx", dtype=str, engine='openpyxl')
            col_kod, col_aciklama = df_nace.columns[0], df_nace.columns[1]
            df_nace[col_kod] = df_nace[col_kod].astype(str).str.strip()
            
            # NACE açıklamasını bul
            found = df_nace[df_nace[col_kod] == nace_kod]
            aciklama = "" if found.empty else str(found.iloc[0][col_aciklama])
            
            replacements["[DEĞİŞTİR:NACEFAALİYET]"] = aciklama
            
            # NACE ve Faaliyet kombinasyonunu oluştur
            if aciklama:
                replacements["[DEĞİŞTİR:NACEVEFAALİYET]"] = f"{nace_kod} - {aciklama}"
            
            logging.info(f"NACE açıklaması bulundu: {nace_kod}")
            
        except Exception as e:
            logging.error(f"NACE açıklama hatası: {e}")
    
    def create_all_documents(self):
        """Tüm belgeleri oluşturur"""
        # 1) Yedek ve belge listesini hazırla
        try:
            replacements, df = self.generator.load_replacements()
            project_name = self.generator.get_project_name(replacements)
        except Exception as e:
            messagebox.showerror("Hata", f"Veri yüklenirken hata oluştu:\n{e}")
            return


        # Yıllık tarih/yıl bilgisi yoksa bugünün tarihiyle doldur (dinamik silme algoritması için)
        tarih_key = "[DEĞİŞTİR:YILLIK:TARİH]"
        yil_key = "[DEĞİŞTİR:YILLIK:YIL]"
        if not replacements.get(tarih_key) or not replacements.get(yil_key):
            today = datetime.datetime.now()
            replacements[tarih_key] = today.strftime("%d.%m.%Y")
            replacements[yil_key] = today.strftime("%Y")

        # RD yöntemini al
        rd_method = replacements.get("[DEĞİŞTİR:RDYONTEMI]", "Matris")
        
        docs = self.generator.get_available_documents(rd_method)
        if not docs:
            messagebox.showwarning("Uyarı", "İşlenecek belge bulunamadı!")
            return

        # 2) İlerleme penceresi oluştur
        prog_win = Toplevel(self.root)
        prog_win.title("Tüm Belgeler Oluşturuluyor…")
        prog_win.geometry("400x100")
        # Pencereyi ekranın tam ortasına yerleştir
        prog_win.update_idletasks()
        width = 400
        height = 100
        x = (prog_win.winfo_screenwidth() // 2) - (width // 2)
        y = (prog_win.winfo_screenheight() // 2) - (height // 2)
        prog_win.geometry(f"{width}x{height}+{x}+{y}")
        tk.Label(prog_win, text="Lütfen bekleyin…", font=(DEFAULT_FONT, 10)).pack(pady=(10, 5))
        pb = ttk.Progressbar(prog_win, orient="horizontal", length=300, mode="determinate")
        pb.pack(pady=(0, 10))
        pb["maximum"] = len(docs)
        pb["value"] = 0

        # 3) Arka planda belge işleme
        def task():
            # PDF oluşturma tercihini al
            self.generator.generate_pdf = self.generate_pdf_var.get()
            # Klasörleri oluştur ve yedeğe veri.xlsx kaydet
            target_folder, backup_folder = self.generator.create_folders(project_name)
            df.to_excel(os.path.join(backup_folder, "veri.xlsx"), index=False, engine='openpyxl')

            success_count = 0
            for idx, doc in enumerate(docs, start=1):
                if self.generator.process_document(doc, replacements,
                        project_name, target_folder, backup_folder):
                    success_count += 1
                # ProgressBar güncelle
                prog_win.after(0, lambda v=idx: pb.config(value=v))

            # İşlem bitince pencereyi kapat
            prog_win.after(0, prog_win.destroy)
            # “Tamamlandı” mesajını ana/root penceresine schedule et
            self.root.after(0, lambda:
                messagebox.showinfo(
                    "Tamamlandı",
                    f"İşlem tamamlandı!\n\n"
                    f"• {success_count}/{len(docs)} belge başarıyla işlendi\n"
                    f"• Belgeler masaüstünde '{project_name}' klasöründe"
                )
            )

        threading.Thread(target=task, daemon=True).start()
    
    def create_selected_documents(self):
        """Belge seçim ekranını açar"""
        try:
            replacements, df = self.generator.load_replacements()
            project_name = self.generator.get_project_name(replacements)
            target_folder, backup_folder = self.generator.create_folders(project_name)
            
            # veri.xlsx'i yedekle
            df.to_excel(os.path.join(backup_folder, "veri.xlsx"), index=False, engine='openpyxl')


            # Yıllık tarih/yıl bilgisi yoksa bugünün tarihiyle doldur (dinamik silme algoritması için)
            tarih_key = "[DEĞİŞTİR:YILLIK:TARİH]"
            yil_key = "[DEĞİŞTİR:YILLIK:YIL]"
            if not replacements.get(tarih_key) or not replacements.get(yil_key):
                today = datetime.datetime.now()
                replacements[tarih_key] = today.strftime("%d.%m.%Y")
                replacements[yil_key] = today.strftime("%Y")

            # RD yöntemi seçimli belge oluşturmada tüm dosyaları göster
            documents = self.generator.get_available_documents()
            if not documents:
                messagebox.showwarning("Uyarı", "İşlenecek belge bulunamadı!")
                return
            
            # Seçim penceresi
            selection_window = Toplevel(self.root)
            selection_window.title("Belge Seç")
            selection_window.geometry("450x550")
            selection_window.configure(bg="#e0e0e0")
            
            tk.Label(selection_window, text="İşlenecek belgeleri seçin:", 
                    font=(DEFAULT_FONT, 10, "bold"),
                    bg="#e0e0e0").pack(pady=10)
            
            # Checkbox frame
            frame = tk.Frame(selection_window, bg="#e0e0e0")
            frame.pack(pady=(0, 10), fill="both", expand=True)
            
            # Canvas ve scrollbar ekle
            canvas = tk.Canvas(frame, bg="#e0e0e0")
            scrollbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview)
            scrollbar.pack(side="right", fill="y")
            canvas.pack(side="left", fill="both", expand=True)
            canvas.configure(yscrollcommand=scrollbar.set)
            
            # Checkbox'ları içeren frame
            checkbox_frame = tk.Frame(canvas, bg="#e0e0e0")
            canvas.create_window((0, 0), window=checkbox_frame, anchor="nw")
            
            # Checkbox'ları oluştur
            checkboxes = {}
            for i, doc in enumerate(documents):
                var = BooleanVar()
                chk = Checkbutton(checkbox_frame, text=doc, variable=var,
                                anchor="w", width=50, bg="#e0e0e0")
                chk.grid(row=i, column=0, sticky="w", padx=10, pady=2)
                checkboxes[doc] = var
            
            # Canvas'ı güncelle
            checkbox_frame.update_idletasks()
            canvas.configure(scrollregion=canvas.bbox("all"))
            
            def process_selected():
                selected_files = [doc for doc, var in checkboxes.items() if var.get()]
                if not selected_files:
                   messagebox.showwarning("Uyarı", "Hiçbir belge seçilmedi!")
                   return

                # Seçim penceresini kapat
                selection_window.destroy()

                # 1) İlerleme penceresi oluştur
                prog_win = Toplevel(self.root)
                prog_win.title("Belgeler Oluşturuluyor…")
                prog_win.geometry("400x100")
                # Pencereyi ekranın tam ortasına yerleştir
                prog_win.update_idletasks()
                width = 400
                height = 100
                x = (prog_win.winfo_screenwidth() // 2) - (width // 2)
                y = (prog_win.winfo_screenheight() // 2) - (height // 2)
                prog_win.geometry(f"{width}x{height}+{x}+{y}")
                tk.Label(prog_win, text="Lütfen bekleyin…", font=(DEFAULT_FONT, 10)).pack(pady=(10, 5))
                pb = ttk.Progressbar(prog_win, orient="horizontal",
                                     length=300, mode="determinate")
                pb.pack(pady=(0,10))
                pb["maximum"] = len(selected_files)
                pb["value"] = 0

                # PDF oluşturma tercihini al
                self.generator.generate_pdf = self.generate_pdf_var.get()
                # 2) Arka planda çalışacak işlev
                def task():
                    success_count = 0
                    for idx, doc in enumerate(selected_files, start=1):
                        ok = self.generator.process_document(
                            doc, replacements, project_name,
                            target_folder, backup_folder
                        )
                        if ok:
                            success_count += 1

                        # ProgressBar'ı ana thread'de güncelle
                        prog_win.after(0, lambda v=idx: pb.config(value=v))

                    # İş bittiğinde pencereleri kapat ve sonucu göster
                    prog_win.after(0, prog_win.destroy)
                    prog_win.after(0, lambda:
                        messagebox.showinfo(
                            "Tamamlandı",
                            f"İşlem tamamlandı!\n\n• {success_count}/{len(selected_files)} belge başarıyla işlendi"
                        )
                    )

                # 3) Thread başlat
                threading.Thread(target=task, daemon=True).start()
            
            # İşlem butonu
            tk.Button(selection_window, text="Seçili Belgeleri Oluştur", 
                     command=process_selected,
                     bg="#4caf50", fg="white", 
                     font=(DEFAULT_FONT, 10, "bold"),
                     width=20).pack(pady=20)
            
        except Exception as e:
            messagebox.showerror("Hata", f"İşlem sırasında hata oluştu:\n{str(e)}")
            logging.error(f"Belge seçim hatası: {e}")
    
    def open_history(self):
        """Evrak geçmişi klasörünü açar"""
        history_path = os.path.join(os.getcwd(), "yedekler")
        os.makedirs(history_path, exist_ok=True)
        
        try:
            if IS_WINDOWS:
                # Windows
                os.startfile(history_path)
            elif IS_MACOS:
                # macOS
                subprocess.run(["open", history_path], check=True)
            else:
                # Linux
                subprocess.run(["xdg-open", history_path], check=True)
        except Exception as e:
            logging.error(f"Klasör açma hatası: {e}")
            messagebox.showinfo("Bilgi", f"Klasör yolu:\n{history_path}")
    
    @staticmethod
    def darken_color(color):
        """Rengi koyulaştırır"""
        return "#0d1235"


# Ana program
if __name__ == "__main__":
    # Gerekli klasörleri kontrol et
    required_folders = ["Evraklar", "yedekler"]
    for folder in required_folders:
        if not os.path.exists(folder):
            os.makedirs(folder)
            logging.info(f"{folder} klasörü oluşturuldu")
    
    # Uygulamayı başlat
    root = tk.Tk()
    app = EvrakGeneratorGUI(root)
    root.mainloop()
