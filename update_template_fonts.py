#!/usr/bin/env python3
"""
update_template_fonts.py

Tüm Evraklar klasöründeki Word ve Excel şablonlarında kullanılan yazı tipini
platformlar arası aynı olacak şekilde DEFAULT_FONT olarak günceller.
"""

import os
import tkinter as tk
import tkinter.font as tkfont

from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook
from openpyxl.styles import Font

# Default font belirleme
_root = tk.Tk()
_root.withdraw()
_available = set(tkfont.families())
_root.destroy()
for _f in ("Arial", "Liberation Sans", "DejaVu Sans", "TkDefaultFont"):
    if _f in _available:
        DEFAULT_FONT = _f
        break
else:
    DEFAULT_FONT = tkfont.nametofont("TkDefaultFont").actual()["family"]

def update_docx(path):
    doc = Document(path)
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = DEFAULT_FONT
            rpr = run._element.rPr
            if rpr is not None:
                rpr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = DEFAULT_FONT
                        rpr = run._element.rPr
                        if rpr is not None:
                            rpr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)
    doc.save(path)

def update_xlsx(path):
    wb = load_workbook(path)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                f = cell.font
                cell.font = Font(
                    name=DEFAULT_FONT,
                    size=f.sz,
                    bold=f.b,
                    italic=f.i,
                    underline=f.u,
                    color=f.color
                )
    wb.save(path)

def main():
    base = os.path.join(os.path.dirname(__file__), 'Evraklar')
    for root, dirs, files in os.walk(base):
        for name in files:
            if name.startswith('~$'):
                continue
            path = os.path.join(root, name)
            if name.lower().endswith('.docx'):
                update_docx(path)
                print(f'Updated DOCX font: {path}')
            elif name.lower().endswith(('.xlsx', '.xlsm', '.xls')):
                update_xlsx(path)
                print(f'Updated XLSX font: {path}')

if __name__ == '__main__':
    main()
